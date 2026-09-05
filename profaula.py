import os
import sys
import re
import json
import threading
import time
import datetime
import urllib.request
import urllib.error
import tkinter as tk
import webbrowser
import subprocess
from tkinter import ttk, filedialog, messagebox, scrolledtext

# Importação defensiva de bibliotecas externas
try:
    import docx
    from docx.shared import Pt, RGBColor, Inches, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_TAB_LEADER
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
except ImportError:
    docx = None

try:
    import pypdf
except ImportError:
    pypdf = None

try:
    from google import genai
    from google.genai import types
except ImportError:
    genai = None
    types = None

# Biblioteca de reconhecimento de voz
try:
    import speech_recognition as sr
except ImportError:
    sr = None

# ==============================================================================
# GERENCIAMENTO DE DIRETÓRIOS, ASSETS E CONFIGURAÇÕES SEGURAS
# ==============================================================================
if getattr(sys, 'frozen', False):
    BUNDLE_DIR = sys._MEIPASS
    BASE_DIR = os.path.dirname(os.path.abspath(sys.executable))
else:
    BUNDLE_DIR = os.path.dirname(os.path.abspath(__file__))
    BASE_DIR = os.path.dirname(os.path.abspath(__file__))

bncc_json_path = os.path.join(BUNDLE_DIR, "bncc_data.json")

DATA_DIR = os.path.join(os.path.expanduser("~"), ".profaula")
os.makedirs(DATA_DIR, exist_ok=True)

# Estrutura de Rascunhos (Gavetas)
DRAFTS_DIR = os.path.join(DATA_DIR, "rascunhos")
os.makedirs(DRAFTS_DIR, exist_ok=True)

API_KEY_FILE = os.path.join(DATA_DIR, ".gemini_api_key.txt")
GROQ_KEY_FILE = os.path.join(DATA_DIR, ".groq_api_key.txt")
OPENROUTER_KEY_FILE = os.path.join(DATA_DIR, ".openrouter_api_key.txt")
CONFIG_DIR_FILE = os.path.join(DATA_DIR, ".profaula_dir.txt")
PROVIDER_CHOICE_FILE = os.path.join(DATA_DIR, ".profaula_provider.txt")
DEV_MODE_FILE = os.path.join(DATA_DIR, ".profaula_devmode.txt")

DISCIPLINAS_LISTA = ["Português", "Matemática", "Ciências", "Geografia", "História", "Arte", "Ensino Religioso", "Educação Física", "Inglês", "Multidisciplinar", "Outra"]

# ==============================================================================
# MÓDULO DE TEMA ESCURO (HELPER GLOBAL)
# ==============================================================================
def apply_theme_to_window(window, is_dark):
    if is_dark:
        bg_color = "#1E1E1E"
        fg_color = "#E0E0E0"
        input_bg = "#2D2D2D"
        input_fg = "#FFFFFF"
        sel_bg = "#005A9E"
    else:
        bg_color = "#F0F0F0"
        fg_color = "#000000"
        input_bg = "#FFFFFF"
        input_fg = "#000000"
        sel_bg = "#0078D7"

    try:
        window.configure(bg=bg_color)
    except: pass

    def update_widgets_recursive(parent):
        for child in parent.winfo_children():
            wclass = child.winfo_class()
            if wclass in ("Text", "ScrolledText", "Entry"):
                try:
                    child.configure(bg=input_bg, fg=input_fg, insertbackground=input_fg, selectbackground=sel_bg)
                except: pass
            elif wclass == "Label":
                try:
                    # Preservar a cor do header original
                    if child.cget("bg") != "#1E293B":
                        current_fg = child.cget("fg")
                        if is_dark and current_fg in ["black", "#000000", "SystemButtonText"]:
                            child.configure(bg=bg_color, fg=fg_color)
                        elif not is_dark and current_fg == fg_color:
                            child.configure(bg=bg_color, fg="black")
                        else:
                            child.configure(bg=bg_color)
                except: pass
            elif wclass == "Frame":
                try:
                    # Preservar cores específicas (Header e botões personalizados se forem frames)
                    if child.cget("bg") not in ["#1E293B", "#2563EB", "#64748B"]:
                        child.configure(bg=bg_color)
                except: pass
            update_widgets_recursive(child)

    update_widgets_recursive(window)

# ==============================================================================
# DIÁLOGO CUSTOMIZADO PARA COPIAR ERRO
# ==============================================================================
def show_error_dialog(parent, title, message, is_dark_mode=False):
    win = tk.Toplevel(parent)
    win.title(title)
    win.geometry("620x480")
    win.minsize(500, 380)

    lbl = tk.Label(win, text="⚠️ Ocorreu uma Falha no Processamento", font=("Segoe UI", 11, "bold"), fg="#DC2626")
    lbl.pack(anchor="w", padx=15, pady=(15, 5))

    txt = scrolledtext.ScrolledText(win, wrap="word", font=("Consolas", 9))
    txt.pack(fill="both", expand=True, padx=15, pady=5)
    txt.insert("1.0", message)
    txt.config(state="disabled")

    btn_frame = tk.Frame(win)
    btn_frame.pack(fill="x", padx=15, pady=12)

    def copy_to_clipboard():
        win.clipboard_clear()
        win.clipboard_append(message)
        messagebox.showinfo("Copiado!", "Log de erro copiado para a área de transferência com sucesso.", parent=win)

    btn_copy = tk.Button(
        btn_frame,
        text="📋 Copiar Erro",
        command=copy_to_clipboard,
        bg="#2563EB",
        fg="white",
        font=("Segoe UI", 9, "bold"),
        padx=12,
        pady=5,
        relief="raised"
    )
    btn_copy.pack(side="left")

    btn_close = tk.Button(
        btn_frame,
        text="Fechar",
        command=win.destroy,
        bg="#64748B",
        fg="white",
        font=("Segoe UI", 9, "bold"),
        padx=12,
        pady=5,
        relief="raised"
    )
    btn_close.pack(side="right")

    apply_theme_to_window(win, is_dark_mode)

# ==============================================================================
# MOTOR MULTI-PROVEDOR SILENCIOSO E RESILIENTE
# ==============================================================================
def call_openai_compatible_api(endpoint_url, api_key, model_name, prompt, system_instruction, temperature=0.6):
    clean_key = api_key.strip()
    if clean_key.lower().startswith("bearer "):
        clean_key = clean_key[7:].strip()

    if not clean_key:
        raise ValueError("Chave de API não informada ou vazia nas configurações.")

    headers = {
        "Authorization": f"Bearer {clean_key}",
        "Content-Type": "application/json",
        "User-Agent": "ProfAula/1.0",
        "HTTP-Referer": "https://github.com/ramonling/profaula",
        "X-Title": "Prof. Aula"
    }

    full_prompt = f"{prompt}\n\nATENÇÃO: Retorne a resposta ESTRITAMENTE em formato JSON válido, sem nenhum texto extra ou marcação."

    payload = {
        "model": model_name,
        "messages": [
            {"role": "system", "content": system_instruction},
            {"role": "user", "content": full_prompt}
        ],
        "temperature": temperature,
        "stream": False,
        "response_format": {"type": "json_object"}
    }

    data_bytes = json.dumps(payload).encode("utf-8")
    req = urllib.request.Request(endpoint_url, data=data_bytes, headers=headers, method="POST")

    try:
        with urllib.request.urlopen(req, timeout=60) as resp:
            res_data = json.loads(resp.read().decode("utf-8"))
    except urllib.error.HTTPError as e:
        try:
            err_body = e.read().decode("utf-8")
        except Exception:
            err_body = "Não foi possível ler o corpo do erro."
        raise Exception(f"HTTP Error {e.code}: {err_body}")

    text = res_data["choices"][0]["message"]["content"]
    usage = res_data.get("usage", {})
    prompt_tokens = usage.get("prompt_tokens", 0)
    completion_tokens = usage.get("completion_tokens", 0)
    total_tokens = usage.get("total_tokens", prompt_tokens + completion_tokens)

    telemetria = {
        "model": f"{model_name}",
        "prompt_tokens": prompt_tokens,
        "response_tokens": completion_tokens,
        "total_tokens": total_tokens
    }
    return text, telemetria

def call_ai_multi_provider(gemini_key, groq_key, openrouter_key, provider_priority, prompt, system_instruction, temperature=0.6, force_fast_model=False):
    providers_queue = []

    if provider_priority == "Groq" and groq_key:
        providers_queue.append("groq")
    elif provider_priority == "OpenRouter (Multi-IA)" and openrouter_key:
        providers_queue.append("openrouter")
    elif provider_priority == "Google Gemini" and gemini_key:
        providers_queue.append("gemini")

    if "gemini" not in providers_queue and gemini_key and genai:
        providers_queue.append("gemini")
    if "groq" not in providers_queue and groq_key:
        providers_queue.append("groq")
    if "openrouter" not in providers_queue and openrouter_key:
        providers_queue.append("openrouter")

    if not providers_queue:
        raise RuntimeError("Nenhuma chave de API válida configurada. Configure ao menos uma chave em ⚙️ Configurações.")

    errors_log = []

    for provider in providers_queue:
        if provider == "gemini" and gemini_key and genai:
            client = genai.Client(api_key=gemini_key.strip())
            gemini_models = ["gemini-3.5-flash-lite", "gemini-3.8-flash"] if force_fast_model else ["gemini-3.8-flash", "gemini-3.7-flash", "gemini-3.5-flash-lite"]
            for model_name in gemini_models:
                try:
                    response = client.models.generate_content(
                        model=model_name,
                        contents=prompt,
                        config=types.GenerateContentConfig(
                            system_instruction=system_instruction,
                            temperature=temperature,
                            response_mime_type="application/json"
                        )
                    )
                    if response and response.text:
                        usage = getattr(response, "usage_metadata", None)
                        p_tok = getattr(usage, "prompt_token_count", 0) if usage else 0
                        r_tok = getattr(usage, "candidates_token_count", 0) if usage else 0
                        t_tok = getattr(usage, "total_token_count", p_tok + r_tok) if usage else 0
                        return response.text, {"model": f"Gemini ({model_name})", "prompt_tokens": p_tok, "response_tokens": r_tok, "total_tokens": t_tok}
                except Exception as e:
                    errors_log.append(f"Gemini [{model_name}]: {e}")

        elif provider == "groq" and groq_key:
            groq_endpoint = "https://api.groq.com/openai/v1/chat/completions"
            groq_models = ["minimaxai/minimax-m2.7"] if force_fast_model else ["openai/gpt-oss-20b", "qwen/qwen3.6-27b", "minimaxai/minimax-m2.7"]
            for model_name in groq_models:
                try:
                    return call_openai_compatible_api(
                        endpoint_url=groq_endpoint,
                        api_key=groq_key.strip(),
                        model_name=model_name,
                        prompt=prompt,
                        system_instruction=system_instruction,
                        temperature=temperature
                    )
                except Exception as e:
                    errors_log.append(f"Groq [{model_name}]: {e}")

        elif provider == "openrouter" and openrouter_key:
            openrouter_models = ["poolside/laguna-xs-2.1:free"] if force_fast_model else ["openrouter/free", "google/gemma-4-31b-it:free", "poolside/laguna-xs-2.1:free", "openai/gpt-oss-120b:free"]
            for model_name in openrouter_models:
                try:
                    return call_openai_compatible_api(
                        endpoint_url="https://openrouter.ai/api/v1/chat/completions",
                        api_key=openrouter_key.strip(),
                        model_name=model_name,
                        prompt=prompt,
                        system_instruction=system_instruction,
                        temperature=temperature
                    )
                except Exception as e:
                    errors_log.append(f"OpenRouter [{model_name}]: {e}")

    raise RuntimeError(f"Todos os provedores e modelos configurados falharam no momento.\n\nDetalhes do Diagnóstico:\n" + "\n".join(errors_log))

# ==============================================================================
# HELPERS DE INTERFACE E TEXTO (SANITIZAÇÃO E ARQUIVOS)
# ==============================================================================
def sanitize_text(text):
    if not text: return ""
    text = str(text).strip(" \"'“’")
    text = re.sub(r"^(\(\s*\)\s*)+", "", text)
    text = re.sub(r"^([A-E]\))\s*([A-E]\))", r"\1", text)
    return text.strip()

def add_context_menu(widget):
    menu = tk.Menu(widget, tearoff=0)
    menu.add_command(label="✂️ Cortar", command=lambda: widget.event_generate("<<Cut>>"))
    menu.add_command(label="📋 Copiar", command=lambda: widget.event_generate("<<Copy>>"))
    menu.add_command(label="📥 Colar", command=lambda: widget.event_generate("<<Paste>>"))
    menu.add_separator()
    menu.add_command(label="🔍 Selecionar Tudo", command=lambda: select_all(widget))
    def show_menu(event): menu.tk_popup(event.x_root, event.y_root)
    widget.bind("<Button-3>", show_menu)
    widget.bind("<Control-a>", lambda e: select_all(widget))
    widget.bind("<Control-A>", lambda e: select_all(widget))

def select_all(widget):
    if isinstance(widget, (tk.Entry, ttk.Entry)):
        widget.select_range(0, tk.END)
        widget.icursor(tk.END)
    elif isinstance(widget, (tk.Text, scrolledtext.ScrolledText)):
        widget.tag_add(tk.SEL, "1.0", tk.END)
        widget.mark_set(tk.INSERT, "1.0")
        widget.see(tk.INSERT)
    return "break"

def paste_clipboard_to_widget(widget):
    try:
        content = widget.clipboard_get()
        if isinstance(widget, (tk.Entry, ttk.Entry)): widget.insert(tk.INSERT, content)
        elif isinstance(widget, (tk.Text, scrolledtext.ScrolledText)): widget.insert(tk.INSERT, content)
    except Exception: pass

def clear_widget(widget):
    if isinstance(widget, (tk.Entry, ttk.Entry)): widget.delete(0, tk.END)
    elif isinstance(widget, (tk.Text, scrolledtext.ScrolledText)): widget.delete("1.0", tk.END)

# ==============================================================================
# CLASSE PRINCIPAL DA APLICAÇÃO GUI
# ==============================================================================
class ProfAulaApp:
    def __init__(self, root):
        self.root = root
        self.APP_VERSION = "1.0.2"
        self.GITHUB_REPO = "ramonling/profaula"

        self.root.title(f"Prof. Aula — Gerador Inteligente de Planos & Atividades v{self.APP_VERSION}")
        self.root.geometry("1020x840")
        self.root.minsize(900, 740)

        self.style = ttk.Style()
        self.style.theme_use("clam")

        self.is_dark_mode = False

        self.current_draft_path = None

        # Variáveis do Cache Fantasma para o Filtro de IA
        self.last_raw_instruction = None
        self.cached_clean_data = None
        self.cached_clean_instruction = None

        self.file_esqueleto_path = tk.StringVar()
        self.file_historico_path = tk.StringVar()
        self.file_livro_path = tk.StringVar()
        self.file_plano_base_path = tk.StringVar()

        self.api_key_var = tk.StringVar(value=self.load_key(API_KEY_FILE))
        self.groq_key_var = tk.StringVar(value=self.load_key(GROQ_KEY_FILE))
        self.openrouter_key_var = tk.StringVar(value=self.load_key(OPENROUTER_KEY_FILE))
        self.save_dir_var = tk.StringVar(value=self.load_save_dir())
        self.provider_priority_var = tk.StringVar(value=self.load_provider_choice())
        self.dev_mode_var = tk.BooleanVar(value=self.load_dev_mode())

        self.last_focused_widget = None
        self.is_recording = False
        self.root.bind_all("<FocusIn>", self.track_focus)

        self.create_header()
        self.create_tabs()
        self.create_footer()

        # Inicia com o tema padronizado ajustado
        self.apply_theme()

    def toggle_theme(self):
        self.is_dark_mode = not self.is_dark_mode
        self.apply_theme()

    def apply_theme(self):
        if self.is_dark_mode:
            bg_color = "#1E1E1E"
            fg_color = "#E0E0E0"
            input_bg = "#2D2D2D"
            input_fg = "#FFFFFF"
            btn_theme_bg = "#EAB308"
            btn_theme_fg = "black"
            btn_theme_text = "☀️ Claro"
            sel_bg = "#005A9E"
        else:
            bg_color = "#F0F0F0"
            fg_color = "#000000"
            input_bg = "#FFFFFF"
            input_fg = "#000000"
            btn_theme_bg = "#475569"
            btn_theme_fg = "white"
            btn_theme_text = "🌙 Escuro"
            sel_bg = "#0078D7"

        self.btn_theme.config(text=btn_theme_text, bg=btn_theme_bg, fg=btn_theme_fg)

        self.style.configure(".", background=bg_color, foreground=fg_color, fieldbackground=input_bg)
        self.style.configure("TFrame", background=bg_color)
        self.style.configure("TLabelframe", background=bg_color, foreground=fg_color)
        self.style.configure("TLabelframe.Label", background=bg_color, foreground=fg_color)
        self.style.configure("TLabel", background=bg_color, foreground=fg_color)
        self.style.configure("TCheckbutton", background=bg_color, foreground=fg_color)
        self.style.map("TCheckbutton", background=[("active", bg_color)])
        self.style.configure("TCombobox", fieldbackground=input_bg, background=bg_color, foreground=input_fg)
        self.style.map("TCombobox", fieldbackground=[("readonly", input_bg)], selectbackground=[("readonly", sel_bg)])
        self.style.configure("TNotebook", background=bg_color)
        self.style.configure("TNotebook.Tab", background=bg_color, foreground=fg_color)
        self.style.map("TNotebook.Tab", background=[("selected", sel_bg)], foreground=[("selected", "#FFFFFF")])
        self.style.configure("Treeview", background=input_bg, foreground=input_fg, fieldbackground=input_bg)
        self.style.map("Treeview", background=[("selected", sel_bg)], foreground=[("selected", "#FFFFFF")])

        apply_theme_to_window(self.root, self.is_dark_mode)

    def track_focus(self, event):
        if isinstance(event.widget, (tk.Entry, ttk.Entry, tk.Text, scrolledtext.ScrolledText)):
            self.last_focused_widget = event.widget

    def create_header(self):
        header_frame = tk.Frame(self.root, bg="#1E293B")
        header_frame.pack(fill="x")

        title_lbl = tk.Label(header_frame, text=f"Prof. Aula — Automação Pedagógica", font=("Segoe UI", 14, "bold"), fg="#F8FAFC", bg="#1E293B")
        title_lbl.pack(side="left", padx=15, pady=15)

        self.btn_mic = tk.Button(header_frame, text="🎤 Ditar Texto", command=self.start_voice_typing, bg="#EF4444", fg="white", font=("Segoe UI", 9, "bold"), relief="flat", padx=10)
        self.btn_mic.pack(side="right", padx=(5, 15), pady=15)

        self.btn_theme = tk.Button(header_frame, text="🌙 Escuro", command=self.toggle_theme, bg="#475569", fg="white", font=("Segoe UI", 9, "bold"), relief="flat", padx=10)
        self.btn_theme.pack(side="right", padx=5, pady=15)

        config_btn = tk.Button(header_frame, text="⚙️ Configurações", command=self.open_settings_dialog, bg="#3B82F6", fg="white", font=("Segoe UI", 9, "bold"), relief="flat", padx=10)
        config_btn.pack(side="right", padx=5, pady=15)

        sobre_btn = tk.Button(header_frame, text="ℹ️ Sobre", command=self.open_sobre_dialog, bg="#8B5CF6", fg="white", font=("Segoe UI", 9, "bold"), relief="flat", padx=10)
        sobre_btn.pack(side="right", padx=5, pady=15)

    def start_voice_typing(self):
        if not sr:
            messagebox.showerror("Erro", "As bibliotecas de áudio não estão instaladas.\nExecute no terminal: pip install SpeechRecognition pyaudio")
            return
        if self.is_recording:
            self.is_recording = False
            self.btn_mic.config(text="⏳ Parando...", bg="#94A3B8")
            return
        if not self.last_focused_widget or not self.last_focused_widget.winfo_exists():
            messagebox.showinfo("Aviso", "Por favor, clique dentro de uma caixa de texto primeiro, e depois clique em Ditar.")
            return

        self.is_recording = True
        self.btn_mic.config(text="⏹️ Parar Gravação", bg="#10B981")
        self.root.update()
        threading.Thread(target=self.process_voice_typing, args=(self.last_focused_widget,), daemon=True).start()

    def process_voice_typing(self, target_widget):
        recognizer = sr.Recognizer()
        try:
            with sr.Microphone() as source:
                recognizer.adjust_for_ambient_noise(source, duration=0.5)
                while self.is_recording:
                    try:
                        audio = recognizer.listen(source, timeout=1, phrase_time_limit=15)
                        if not self.is_recording: break
                        texto_reconhecido = recognizer.recognize_google(audio, language="pt-BR")
                        self.root.after(0, self.insert_text_to_widget, target_widget, texto_reconhecido)
                    except sr.WaitTimeoutError: continue
                    except sr.UnknownValueError: continue
                    except Exception as e:
                        print(f"Erro leve no áudio: {e}")
                        time.sleep(1)
        except Exception as e:
            self.root.after(0, messagebox.showerror, "Erro de Microfone", f"Não foi possível iniciar o microfone: {e}")
        finally:
            self.is_recording = False
            self.root.after(0, lambda: self.btn_mic.config(text="🎤 Ditar Texto", bg="#EF4444"))

    def insert_text_to_widget(self, widget, text):
        if not text: return
        text = text.capitalize() + " "
        try:
            if isinstance(widget, (tk.Entry, ttk.Entry)): widget.insert(tk.INSERT, text)
            elif isinstance(widget, (tk.Text, scrolledtext.ScrolledText)): widget.insert(tk.INSERT, text)
        except Exception: pass

    # Funções de IO Básicas
    def load_key(self, filepath):
        if os.path.exists(filepath):
            try:
                with open(filepath, "r", encoding="utf-8") as f: return f.read().strip()
            except Exception: return ""
        return ""

    def load_provider_choice(self):
        if os.path.exists(PROVIDER_CHOICE_FILE):
            try:
                with open(PROVIDER_CHOICE_FILE, "r", encoding="utf-8") as f: return f.read().strip()
            except Exception: pass
        return "Auto (Fallback Inteligente)"

    def load_save_dir(self):
        if os.path.exists(CONFIG_DIR_FILE):
            try:
                with open(CONFIG_DIR_FILE, "r", encoding="utf-8") as f:
                    path = f.read().strip()
                    if os.path.isdir(path): return path
            except Exception: pass
        docs_dir = os.path.join(os.path.expanduser("~"), "Documents")
        if os.path.isdir(docs_dir): return docs_dir
        return BASE_DIR

    def load_dev_mode(self):
        if os.path.exists(DEV_MODE_FILE):
            try:
                with open(DEV_MODE_FILE, "r", encoding="utf-8") as f:
                    return f.read().strip() == "True"
            except Exception: pass
        return False

    def abrir_link(self, url):
        # 1. Sistema Anti-Duplo Clique (Debounce de 1 segundo)
        if hasattr(self, '_last_link_click') and time.time() - self._last_link_click < 1.0:
            return
        self._last_link_click = time.time()

        # 2. Limpeza Brutal de Markdown (Remove os colchetes inseridos ao copiar o código)
        clean_url = str(url).strip()
        clean_url = re.sub(r'\[.*?\]\(|\)|\[|\]|<|>', '', clean_url)

        try:
            # Tenta abrir pelo navegador nativo do Python
            sucesso = webbrowser.open_new_tab(clean_url)
            if sucesso:
                return # Se abriu com sucesso, PARA a função aqui para não abrir duplicado!
        except Exception as e:
            print(f"Webbrowser falhou ({e}), tentando fallback do Sistema Operacional...")

        # 3. Fallback do Sistema Operacional (Só roda se o de cima falhar)
        try:
            if sys.platform == 'win32':
                os.startfile(clean_url)
            elif sys.platform == 'darwin':
                subprocess.Popen(['open', clean_url])
            else:
                subprocess.Popen(['xdg-open', clean_url], stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL)
        except Exception as e:
            print(f"Erro extremo ao abrir link no SO: {e}")

    # ==============================================================================
    # FUNÇÕES DE EXTRAÇÃO DE ARQUIVO E PARSER
    # ==============================================================================
    def pick_file(self, target_var, file_types):
        path = filedialog.askopenfilename(filetypes=file_types)
        if path:
            target_var.set(path)

    def extract_text_from_file(self, filepath):
        if not filepath or not os.path.exists(filepath):
            return ""
        ext = os.path.splitext(filepath)[1].lower()
        text = ""
        try:
            if ext == ".txt":
                with open(filepath, "r", encoding="utf-8", errors="ignore") as f:
                    text = f.read()
            elif ext == ".docx" and docx:
                doc = docx.Document(filepath)
                text = "\n".join([p.text for p in doc.paragraphs if p.text])
            elif ext == ".pdf" and pypdf:
                reader = pypdf.PdfReader(filepath)
                for page in reader.pages:
                    extracted = page.extract_text()
                    if extracted:
                        text += extracted + "\n"
        except Exception as e:
            print(f"Erro ao ler arquivo {filepath}: {e}")
        return text.strip()

    def parse_ai_json(self, response_text):
        try:
            clean_text = response_text.strip()
            if clean_text.startswith("```json"):
                clean_text = clean_text[7:]
            if clean_text.startswith("```"):
                clean_text = clean_text[3:]
            if clean_text.endswith("```"):
                clean_text = clean_text[:-3]

            return json.loads(clean_text.strip())
        except json.JSONDecodeError as e:
            print(f"Falha ao decodificar JSON da IA: {e}")
            print(f"Resposta bruta recebida: {response_text}")
            raise ValueError("A IA não retornou um formato JSON válido.")

    # Dialogs Auxiliares
    def open_sobre_dialog(self):
        win = tk.Toplevel(self.root)
        win.title("Sobre o Prof. Aula")
        win.geometry("400x480")
        win.resizable(False, False)
        tk.Label(win, text="Prof. Aula", font=("Segoe UI", 18, "bold")).pack(pady=(25, 2))
        tk.Label(win, text=f"Versão {self.APP_VERSION}", font=("Segoe UI", 10)).pack(pady=(0, 20))
        tk.Label(win, text="Desenvolvido por:", font=("Segoe UI", 9)).pack()
        tk.Label(win, text="Ramon", font=("Segoe UI", 12, "bold")).pack(pady=(0, 10))
        btn_frame = tk.Frame(win)
        btn_frame.pack(pady=10)

        tk.Button(btn_frame, text="📸 Instagram (@ramonchvr)", command=lambda: self.abrir_link("[https://instagram.com/ramonchvr](https://instagram.com/ramonchvr)"), width=30, bg="#F43F5E", fg="white", font=("Segoe UI", 9, "bold")).pack(pady=5)
        tk.Button(btn_frame, text="🐙 Acessar GitHub Oficial", command=lambda: self.abrir_link(f"[https://github.com/](https://github.com/){self.GITHUB_REPO}"), width=30, bg="#334155", fg="white", font=("Segoe UI", 9, "bold")).pack(pady=5)
        tk.Button(btn_frame, text="🐛 Relatar Bug / Enviar Feedback", command=lambda: self.abrir_link("mailto:profaulaai@gmail.com?subject=Feedback/Bug%20-%20Prof.%20Aula"), width=30, bg="#0EA5E9", fg="white", font=("Segoe UI", 9, "bold")).pack(pady=5)

        ttk.Separator(win, orient='horizontal').pack(fill='x', padx=20, pady=15)
        self.lbl_update_status = tk.Label(win, text="", font=("Segoe UI", 9), fg="#475569")
        self.lbl_update_status.pack(pady=(5, 5))
        tk.Button(win, text="🔄 Verificar Atualizações", command=self.check_for_updates, bg="#10B981", fg="white", font=("Segoe UI", 10, "bold"), padx=15).pack(pady=5)
        apply_theme_to_window(win, self.is_dark_mode)

    def check_for_updates(self):
        self.lbl_update_status.config(text="Verificando atualizações no GitHub...")
        self.root.update()
        def _check():
            try:
                import ssl
                ctx = ssl.create_default_context()
                ctx.check_hostname = False
                ctx.verify_mode = ssl.CERT_NONE

                # Higieniza a URL da API do GitHub para garantir que não tenha artefatos
                url = f"[https://api.github.com/repos/](https://api.github.com/repos/){self.GITHUB_REPO}/releases/latest"
                clean_url = re.sub(r'\[.*?\]\(|\)|\[|\]|<|>', '', url).strip()

                req = urllib.request.Request(clean_url, headers={"User-Agent": "ProfAulaApp/1.0"})

                with urllib.request.urlopen(req, timeout=10, context=ctx) as response:
                    data = json.loads(response.read().decode("utf-8"))
                    latest_version = data.get("tag_name", "").lstrip("v")
                    current_version_num = self.APP_VERSION.lstrip("v")
                    is_newer = False
                    if latest_version:
                        try:
                            latest_parts = [int(x) for x in latest_version.split('.')]
                            current_parts = [int(x) for x in current_version_num.split('.')]
                            is_newer = latest_parts > current_parts
                        except Exception:
                            is_newer = latest_version > current_version_num
                    if is_newer:
                        self.root.after(0, lambda: self.lbl_update_status.config(text=f"Nova versão disponível: v{latest_version}!"))
                        self.root.after(0, lambda: messagebox.showinfo("Atualização Disponível", f"Uma nova versão (v{latest_version}) está disponível!\nVocê será redirecionado para download.", parent=self.root))
                        self.root.after(0, lambda: self.abrir_link(f"[https://github.com/](https://github.com/){self.GITHUB_REPO}/releases/latest"))
                    else:
                        self.root.after(0, lambda: self.lbl_update_status.config(text="Você já está na versão mais recente."))

            except urllib.error.HTTPError as e:
                if e.code == 404:
                    msg = "Nenhuma atualização (O repositório não possui Releases)."
                elif e.code == 403:
                    msg = "Limite de consultas do GitHub atingido. Tente mais tarde."
                else:
                    msg = f"Erro HTTP {e.code} ao verificar versão."
                self.root.after(0, lambda: self.lbl_update_status.config(text=msg))

            except Exception as e:
                erro_txt = str(e).split(']')[-1].strip()[:40] if ']' in str(e) else str(e)[:40]
                self.root.after(0, lambda: self.lbl_update_status.config(text=f"Erro de Rede: {erro_txt}"))

        threading.Thread(target=_check, daemon=True).start()

    def open_settings_dialog(self):
        win = tk.Toplevel(self.root)
        win.title("Configurações Multi-IA do Prof. Aula")
        win.geometry("580x480")
        win.resizable(False, False)

        tk.Label(win, text="Chave Google Gemini API:", font=("Segoe UI", 9, "bold")).pack(anchor="w", padx=20, pady=(12, 2))
        k1 = ttk.Entry(win, textvariable=self.api_key_var, width=65, show="*")
        k1.pack(anchor="w", padx=20, pady=(0, 5))
        add_context_menu(k1)

        tk.Label(win, text="Chave Groq API (gsk_...):", font=("Segoe UI", 9, "bold")).pack(anchor="w", padx=20, pady=(5, 2))
        k2 = ttk.Entry(win, textvariable=self.groq_key_var, width=65, show="*")
        k2.pack(anchor="w", padx=20, pady=(0, 5))
        add_context_menu(k2)

        tk.Label(win, text="Chave OpenRouter API (Opcional):", font=("Segoe UI", 9, "bold")).pack(anchor="w", padx=20, pady=(5, 2))
        k3 = ttk.Entry(win, textvariable=self.openrouter_key_var, width=65, show="*")
        k3.pack(anchor="w", padx=20, pady=(0, 10))
        add_context_menu(k3)

        tk.Label(win, text="Provedor de IA Prioritário:", font=("Segoe UI", 9, "bold")).pack(anchor="w", padx=20, pady=(2, 2))
        cb_prov = ttk.Combobox(win, textvariable=self.provider_priority_var, values=["Auto (Fallback Inteligente)", "Google Gemini", "Groq", "OpenRouter (Multi-IA)"], state="readonly", width=35)
        cb_prov.pack(anchor="w", padx=20, pady=(0, 10))

        tk.Label(win, text="Pasta Padrão para Salvar Documentos (.docx):", font=("Segoe UI", 9, "bold")).pack(anchor="w", padx=20, pady=(5, 2))
        dir_box = ttk.Frame(win)
        dir_box.pack(fill="x", padx=20, pady=(0, 10))
        dir_entry = ttk.Entry(dir_box, textvariable=self.save_dir_var, width=50)
        dir_entry.pack(side="left", fill="x", expand=True, padx=(0, 5))
        add_context_menu(dir_entry)
        ttk.Button(dir_box, text="Procurar...", command=lambda: self.save_dir_var.set(filedialog.askdirectory(initialdir=self.save_dir_var.get()) or self.save_dir_var.get())).pack(side="right")

        ttk.Separator(win, orient='horizontal').pack(fill='x', padx=20, pady=10)

        # Checkbox do Modo Desenvolvedor
        cb_dev = ttk.Checkbutton(win, text="🔍 Modo Transparência (Raio-X da IA)", variable=self.dev_mode_var)
        cb_dev.pack(anchor="w", padx=20, pady=(0, 2))
        tk.Label(win, text="Exibe os parâmetros técnicos que serão enviados ao gerador na tela de Prévia.", font=("Segoe UI", 8), fg="#64748B").pack(anchor="w", padx=35, pady=(0, 15))

        def save_all():
            with open(API_KEY_FILE, "w", encoding="utf-8") as f: f.write(self.api_key_var.get().strip())
            with open(GROQ_KEY_FILE, "w", encoding="utf-8") as f: f.write(self.groq_key_var.get().strip())
            with open(OPENROUTER_KEY_FILE, "w", encoding="utf-8") as f: f.write(self.openrouter_key_var.get().strip())
            with open(PROVIDER_CHOICE_FILE, "w", encoding="utf-8") as f: f.write(self.provider_priority_var.get().strip())
            with open(DEV_MODE_FILE, "w", encoding="utf-8") as f: f.write(str(self.dev_mode_var.get()))

            directory = self.save_dir_var.get().strip()
            if not directory or not os.path.isdir(directory): directory = self.load_save_dir()
            with open(CONFIG_DIR_FILE, "w", encoding="utf-8") as f: f.write(directory)

            messagebox.showinfo("Sucesso", "Configurações de IA salvas com sucesso!", parent=win)
            win.destroy()

        tk.Button(win, text="Salvar Configurações", command=save_all, bg="#10B981", fg="white", font=("Segoe UI", 10, "bold"), padx=15).pack(pady=5)
        apply_theme_to_window(win, self.is_dark_mode)

    # ==============================================================================
    # GERENCIADOR DE RASCUNHOS (AS GAVETAS DO PROFESSOR)
    # ==============================================================================
    def get_state_dict(self):
        data = {
            "file_esqueleto": self.file_esqueleto_path.get(),
            "esqueleto": self.txt_esqueleto.get("1.0", tk.END).strip(),
            "file_historico": self.file_historico_path.get(),
            "historico": self.txt_historico.get("1.0", tk.END).strip(),
            "file_livro": self.file_livro_path.get(),
            "livro": self.txt_livro.get("1.0", tk.END).strip(),
            "file_plano_base": self.file_plano_base_path.get(),
            "instrucoes_atividade": self.txt_instrucoes_atividade.get("1.0", tk.END).strip(),
            "genero": self.ent_genero.get(),
            "extra": self.ent_extra.get(),
            "frase": self.combo_frases.get(),
            "observacoes": self.txt_observacoes.get("1.0", tk.END).strip(),
            "ano": self.combo_ano.get(),
            "duracao": self.combo_duracao.get(),
            "multi": self.var_multi.get(),
            "auto_bncc": self.var_auto_bncc.get(),
            "bncc_manual": self.ent_bncc_manual.get(),
            "escola_nome": self.ent_escola_nome.get(),
            "perfil_turma": self.txt_perfil_turma.get("1.0", tk.END).strip(),
            "infraestrutura": self.txt_infraestrutura.get("1.0", tk.END).strip(),
            "realidade_local": self.txt_realidade_local.get("1.0", tk.END).strip(),
            "res_lousa_caderno": self.var_res_lousa_caderno.get(),
            "res_sem_impressao": self.var_res_sem_impressao.get(),
            "res_datashow": self.var_res_datashow.get(),
            "res_patio": self.var_res_patio.get(),
            "res_laboratorio": self.var_res_laboratorio.get()
        }
        if hasattr(self, 'var_gabarito'):
            data["cfg_gabarito"] = self.var_gabarito.get()
            data["cfg_simplificar"] = self.var_simplificar.get()
            data["cfg_texto_apoio"] = self.var_texto_apoio.get()
            data["cfg_tamanho_texto"] = self.combo_tamanho_texto.get()
            data["cfg_nivel"] = self.combo_nivel.get()
            data["cfg_foco"] = self.combo_foco.get()
            data["cfg_formato"] = self.combo_formato.get()
            data["cfg_filtro_modo"] = self.combo_filtro_modo.get()

        if hasattr(self, 'combo_margem_plano'):
            data["cfg_margem_plano"] = self.combo_margem_plano.get()
            data["cfg_coluna_plano"] = self.combo_coluna_plano.get()
            data["cfg_fonte_plano"] = self.combo_fonte_plano.get()
            data["cfg_entrelinhas_plano"] = self.combo_entrelinhas_plano.get()
            data["cfg_espaco_plano"] = self.combo_espaco_plano.get()

        if hasattr(self, 'combo_margem'):
            data["cfg_margem"] = self.combo_margem.get()
            data["cfg_coluna"] = self.combo_coluna.get()
            data["cfg_fonte"] = self.combo_fonte.get()
            data["cfg_entrelinhas"] = self.combo_entrelinhas.get()
            data["cfg_espaco"] = self.combo_espaco.get()
        return data

    def set_state_dict(self, data):
        if "file_esqueleto" in data: self.file_esqueleto_path.set(data["file_esqueleto"])
        if "esqueleto" in data:
            self.txt_esqueleto.delete("1.0", tk.END)
            self.txt_esqueleto.insert("1.0", data["esqueleto"])
        if "file_historico" in data: self.file_historico_path.set(data["file_historico"])
        if "historico" in data:
            self.txt_historico.delete("1.0", tk.END)
            self.txt_historico.insert("1.0", data["historico"])
        if "file_livro" in data: self.file_livro_path.set(data["file_livro"])
        if "livro" in data:
            self.txt_livro.delete("1.0", tk.END)
            self.txt_livro.insert("1.0", data["livro"])
        if "file_plano_base" in data: self.file_plano_base_path.set(data["file_plano_base"])
        if hasattr(self, 'txt_instrucoes_atividade') and "instrucoes_atividade" in data:
            self.txt_instrucoes_atividade.delete("1.0", tk.END)
            self.txt_instrucoes_atividade.insert("1.0", data["instrucoes_atividade"])
        if "genero" in data:
            self.ent_genero.delete(0, tk.END)
            self.ent_genero.insert(0, data["genero"])
        if "extra" in data:
            self.ent_extra.delete(0, tk.END)
            self.ent_extra.insert(0, data["extra"])
        if "frase" in data and data["frase"] in self.frases_list: self.combo_frases.set(data["frase"])
        if "observacoes" in data:
            self.txt_observacoes.delete("1.0", tk.END)
            self.txt_observacoes.insert("1.0", data["observacoes"])
        if "ano" in data and data["ano"] in self.combo_ano['values']: self.combo_ano.set(data["ano"])
        if "duracao" in data: self.combo_duracao.set(data["duracao"])
        if "multi" in data: self.var_multi.set(data["multi"])
        if "auto_bncc" in data: self.var_auto_bncc.set(data["auto_bncc"])
        if "bncc_manual" in data:
            self.ent_bncc_manual.delete(0, tk.END)
            self.ent_bncc_manual.insert(0, data["bncc_manual"])
        if "escola_nome" in data:
            self.ent_escola_nome.delete(0, tk.END)
            self.ent_escola_nome.insert(0, data["escola_nome"])
        if "perfil_turma" in data:
            self.txt_perfil_turma.delete("1.0", tk.END)
            self.txt_perfil_turma.insert("1.0", data["perfil_turma"])
        if "infraestrutura" in data:
            self.txt_infraestrutura.delete("1.0", tk.END)
            self.txt_infraestrutura.insert("1.0", data["infraestrutura"])
        if "realidade_local" in data:
            self.txt_realidade_local.delete("1.0", tk.END)
            self.txt_realidade_local.insert("1.0", data["realidade_local"])
        if "res_lousa_caderno" in data: self.var_res_lousa_caderno.set(data["res_lousa_caderno"])
        if "res_sem_impressao" in data: self.var_res_sem_impressao.set(data["res_sem_impressao"])
        if "res_datashow" in data: self.var_res_datashow.set(data["res_datashow"])
        if "res_patio" in data: self.var_res_patio.set(data["res_patio"])
        if "res_laboratorio" in data: self.var_res_laboratorio.set(data["res_laboratorio"])
        if hasattr(self, 'var_gabarito'):
            if "cfg_gabarito" in data: self.var_gabarito.set(data["cfg_gabarito"])
            if "cfg_simplificar" in data: self.var_simplificar.set(data["cfg_simplificar"])
            if "cfg_texto_apoio" in data: self.var_texto_apoio.set(data["cfg_texto_apoio"])
            if "cfg_tamanho_texto" in data and data["cfg_tamanho_texto"] in self.combo_tamanho_texto['values']: self.combo_tamanho_texto.set(data["cfg_tamanho_texto"])
            if "cfg_nivel" in data and data["cfg_nivel"] in self.combo_nivel['values']: self.combo_nivel.set(data["cfg_nivel"])
            if "cfg_foco" in data and data["cfg_foco"] in self.combo_foco['values']: self.combo_foco.set(data["cfg_foco"])
            if "cfg_formato" in data and data["cfg_formato"] in self.combo_formato['values']: self.combo_formato.set(data["cfg_formato"])
            if "cfg_filtro_modo" in data and data["cfg_filtro_modo"] in self.combo_filtro_modo['values']: self.combo_filtro_modo.set(data["cfg_filtro_modo"])

        if hasattr(self, 'combo_margem_plano'):
            if "cfg_margem_plano" in data and data["cfg_margem_plano"] in self.combo_margem_plano['values']: self.combo_margem_plano.set(data["cfg_margem_plano"])
            if "cfg_coluna_plano" in data and data["cfg_coluna_plano"] in self.combo_coluna_plano['values']: self.combo_coluna_plano.set(data["cfg_coluna_plano"])
            if "cfg_fonte_plano" in data and data["cfg_fonte_plano"] in self.combo_fonte_plano['values']: self.combo_fonte_plano.set(data["cfg_fonte_plano"])
            if "cfg_entrelinhas_plano" in data and data["cfg_entrelinhas_plano"] in self.combo_entrelinhas_plano['values']: self.combo_entrelinhas_plano.set(data["cfg_entrelinhas_plano"])
            if "cfg_espaco_plano" in data and data["cfg_espaco_plano"] in self.combo_espaco_plano['values']: self.combo_espaco_plano.set(data["cfg_espaco_plano"])

        if hasattr(self, 'combo_margem'):
            if "cfg_margem" in data and data["cfg_margem"] in self.combo_margem['values']: self.combo_margem.set(data["cfg_margem"])
            if "cfg_coluna" in data and data["cfg_coluna"] in self.combo_coluna['values']: self.combo_coluna.set(data["cfg_coluna"])
            if "cfg_fonte" in data and data["cfg_fonte"] in self.combo_fonte['values']: self.combo_fonte.set(data["cfg_fonte"])
            if "cfg_entrelinhas" in data and data["cfg_entrelinhas"] in self.combo_entrelinhas['values']: self.combo_entrelinhas.set(data["cfg_entrelinhas"])
            if "cfg_espaco" in data and data["cfg_espaco"] in self.combo_espaco['values']: self.combo_espaco.set(data["cfg_espaco"])

    def quick_save_draft(self):
        if self.current_draft_path and os.path.exists(self.current_draft_path):
            try:
                with open(self.current_draft_path, "w", encoding="utf-8") as f:
                    json.dump(self.get_state_dict(), f, ensure_ascii=False, indent=4)
                self.lbl_status.config(text="✅ Rascunho atualizado com sucesso!")
            except Exception as e:
                messagebox.showerror("Erro", f"Falha ao salvar: {e}")
        else:
            self.open_save_as_dialog()

    def open_save_as_dialog(self):
        win = tk.Toplevel(self.root)
        win.title("Salvar Rascunho Como...")
        win.geometry("450x250")
        win.resizable(False, False)

        tk.Label(win, text="Selecione a Gaveta (Disciplina):", font=("Segoe UI", 10, "bold")).pack(anchor="w", padx=20, pady=(15, 2))
        combo_disc = ttk.Combobox(win, values=DISCIPLINAS_LISTA, width=45)
        combo_disc.current(0)
        combo_disc.pack(anchor="w", padx=20, pady=(0, 15))

        tk.Label(win, text="Nome do Rascunho (Opcional):", font=("Segoe UI", 10, "bold")).pack(anchor="w", padx=20, pady=(0, 2))

        sugestao_nome = f"Rascunho_{self.combo_ano.get().replace(' ', '')}_{datetime.datetime.now().strftime('%d-%m-%Y')}"
        ent_nome = ttk.Entry(win, width=48)
        ent_nome.insert(0, sugestao_nome)
        ent_nome.pack(anchor="w", padx=20, pady=(0, 20))
        add_context_menu(ent_nome)

        def salvar_agora():
            disc = combo_disc.get().strip().replace("/", "_").replace("\\", "_")
            if not disc: disc = "Outra"
            nome = ent_nome.get().strip().replace("/", "_").replace("\\", "_")
            if not nome: nome = sugestao_nome

            disc_dir = os.path.join(DRAFTS_DIR, disc)
            os.makedirs(disc_dir, exist_ok=True)

            filename = f"{nome}.json"
            filepath = os.path.join(disc_dir, filename)

            try:
                with open(filepath, "w", encoding="utf-8") as f:
                    json.dump(self.get_state_dict(), f, ensure_ascii=False, indent=4)
                self.current_draft_path = filepath
                self.lbl_status.config(text=f"✅ Salvo em: {disc} / {nome}")
                win.destroy()
            except Exception as e:
                messagebox.showerror("Erro", f"Falha ao salvar rascunho:\n{e}", parent=win)

        tk.Button(win, text="💾 Guardar na Gaveta", command=salvar_agora, bg="#059669", fg="white", font=("Segoe UI", 10, "bold"), padx=15).pack(pady=5)
        apply_theme_to_window(win, self.is_dark_mode)

    def open_draft_manager(self):
        win = tk.Toplevel(self.root)
        win.title("📚 Meus Rascunhos (Fichário)")
        win.geometry("600x450")
        win.minsize(500, 400)

        tk.Label(win, text="Selecione um Rascunho Antigo para Carregar:", font=("Segoe UI", 11, "bold")).pack(anchor="w", padx=15, pady=(15, 5))

        tree_frame = ttk.Frame(win)
        tree_frame.pack(fill="both", expand=True, padx=15, pady=5)

        tree_scroll = ttk.Scrollbar(tree_frame)
        tree_scroll.pack(side="right", fill="y")

        self.tree = ttk.Treeview(tree_frame, yscrollcommand=tree_scroll.set, selectmode="browse")
        self.tree.pack(side="left", fill="both", expand=True)
        tree_scroll.config(command=self.tree.yview)

        self.populate_treeview()

        btn_frame = ttk.Frame(win)
        btn_frame.pack(fill="x", padx=15, pady=15)

        def carregar_rascunho_selecionado():
            selected_item = self.tree.selection()
            if not selected_item: return
            item_id = selected_item[0]

            item_values = self.tree.item(item_id, "values")
            if not item_values: return

            filepath = item_values[0]
            if os.path.exists(filepath):
                try:
                    with open(filepath, "r", encoding="utf-8") as f:
                        data = json.load(f)
                    self.set_state_dict(data)
                    self.current_draft_path = filepath
                    self.lbl_status.config(text="✅ Rascunho carregado com sucesso!")
                    win.destroy()
                except Exception as e:
                    messagebox.showerror("Erro", f"Falha ao ler o arquivo:\n{e}", parent=win)

        def excluir_rascunho():
            selected_item = self.tree.selection()
            if not selected_item: return
            item_id = selected_item[0]
            item_values = self.tree.item(item_id, "values")
            if not item_values: return

            filepath = item_values[0]
            if messagebox.askyesno("Confirmar", "Tem certeza que deseja apagar este rascunho permanentemente?", parent=win):
                try:
                    os.remove(filepath)
                    self.tree.delete(item_id)
                    if self.current_draft_path == filepath: self.current_draft_path = None
                except Exception as e:
                    messagebox.showerror("Erro", f"Falha ao apagar: {e}", parent=win)

        tk.Button(btn_frame, text="📥 Carregar Rascunho", command=carregar_rascunho_selecionado, bg="#2563EB", fg="white", font=("Segoe UI", 10, "bold"), padx=15).pack(side="right", padx=5)
        tk.Button(btn_frame, text="🗑️ Apagar", command=excluir_rascunho, bg="#EF4444", fg="white", font=("Segoe UI", 9, "bold"), padx=10).pack(side="left", padx=5)
        apply_theme_to_window(win, self.is_dark_mode)

    def populate_treeview(self):
        for i in self.tree.get_children(): self.tree.delete(i)

        for disciplina in sorted(os.listdir(DRAFTS_DIR)):
            disc_path = os.path.join(DRAFTS_DIR, disciplina)
            if os.path.isdir(disc_path):
                arquivos = [f for f in os.listdir(disc_path) if f.endswith(".json")]
                if arquivos:
                    folder_id = self.tree.insert("", "end", text=f"📂 {disciplina}", open=True)
                    for arq in sorted(arquivos, reverse=True):
                        arq_path = os.path.join(disc_path, arq)
                        nome_limpo = arq.replace(".json", "")
                        self.tree.insert(folder_id, "end", text=f"📄 {nome_limpo}", values=(arq_path,))

    def load_state(self):
        if os.path.exists(STATE_FILE):
            try:
                with open(STATE_FILE, "r", encoding="utf-8") as f:
                    data = json.load(f)
                self.set_state_dict(data)
                self.current_draft_path = STATE_FILE
            except Exception as e: print(f"Aviso ao carregar rascunho anterior: {e}")

    # ==============================================================================
    # CRIAÇÃO DAS ABAS
    # ==============================================================================
    def create_tabs(self):
        self.notebook = ttk.Notebook(self.root)
        self.notebook.pack(fill="both", expand=True, padx=10, pady=10)

        self.tab1 = ttk.Frame(self.notebook)
        self.tab2 = ttk.Frame(self.notebook)
        self.tab3 = ttk.Frame(self.notebook)
        self.tab4 = ttk.Frame(self.notebook)
        self.tab5 = ttk.Frame(self.notebook)
        self.tab6 = ttk.Frame(self.notebook)
        self.tab7 = ttk.Frame(self.notebook)

        self.notebook.add(self.tab1, text=" 📝 1. Esqueleto ")
        self.notebook.add(self.tab2, text=" 📚 2. Histórico ")
        self.notebook.add(self.tab3, text=" 📖 3. Livro (PDF) ")
        self.notebook.add(self.tab4, text=" 💡 4. Diretrizes ")
        self.notebook.add(self.tab5, text=" 🎯 5. BNCC ")
        self.notebook.add(self.tab6, text=" 🏫 6. Contexto ")
        self.notebook.add(self.tab7, text=" 🖨️ 7. Atividades ")

        self.setup_tab1()
        self.setup_tab2()
        self.setup_tab3()
        self.setup_tab4()
        self.setup_tab5()
        self.setup_tab6()
        self.setup_tab7()

    def create_field_toolbar(self, parent, target_widget):
        bar = ttk.Frame(parent)
        bar.pack(anchor="e", pady=(0, 2))
        ttk.Button(bar, text="📋 Colar", width=8, command=lambda: paste_clipboard_to_widget(target_widget)).pack(side="left", padx=2)
        ttk.Button(bar, text="❌ Limpar", width=8, command=lambda: clear_widget(target_widget)).pack(side="left", padx=2)

    def setup_tab1(self):
        frame = ttk.LabelFrame(self.tab1, text=" Modelo Estrutural da Escola (.docx / .txt) ", padding=15)
        frame.pack(fill="both", expand=True, padx=15, pady=15)

        ttk.Label(frame, text="Selecione o arquivo do Word ou TXT que contém a estrutura oficial exigida pela sua escola:").pack(anchor="w", pady=(0, 5))
        file_box = ttk.Frame(frame)
        file_box.pack(fill="x", pady=5)
        ent = ttk.Entry(file_box, textvariable=self.file_esqueleto_path, width=60)
        ent.pack(side="left", fill="x", expand=True, padx=(0, 5))
        add_context_menu(ent)
        ttk.Button(file_box, text="Procurar...", command=lambda: self.pick_file(self.file_esqueleto_path, [("Word / Text", "*.docx *.txt")])).pack(side="right")

        ttk.Label(frame, text="Ou digite/cole a estrutura do plano diretamente no campo abaixo:").pack(anchor="w", pady=(15, 5))
        self.txt_esqueleto = scrolledtext.ScrolledText(frame, height=9, font=("Consolas", 10))
        self.txt_esqueleto.pack(fill="both", expand=True)
        add_context_menu(self.txt_esqueleto)
        self.create_field_toolbar(frame, self.txt_esqueleto)

        frame_cfg_plano = ttk.LabelFrame(frame, text=" Configurações de Impressão e Layout (Apenas para o Plano de Aula) ", padding=10)
        frame_cfg_plano.pack(fill="x", pady=(10, 0))
        ttk.Label(frame_cfg_plano, text="Margens:").grid(row=0, column=0, padx=5, pady=5, sticky="e")
        self.combo_margem_plano = ttk.Combobox(frame_cfg_plano, values=["Normal", "Estreita", "Moderada", "Larga"], state="readonly", width=15)
        self.combo_margem_plano.current(0)
        self.combo_margem_plano.grid(row=0, column=1, padx=5, pady=5, sticky="w")
        ttk.Label(frame_cfg_plano, text="Colunas:").grid(row=0, column=2, padx=5, pady=5, sticky="e")
        self.combo_coluna_plano = ttk.Combobox(frame_cfg_plano, values=["1 Coluna (Padrão)", "2 Colunas (Apostila)"], state="readonly", width=20)
        self.combo_coluna_plano.current(0)
        self.combo_coluna_plano.grid(row=0, column=3, padx=5, pady=5, sticky="w")
        ttk.Label(frame_cfg_plano, text="Tamanho da Fonte:").grid(row=0, column=4, padx=5, pady=5, sticky="e")
        self.combo_fonte_plano = ttk.Combobox(frame_cfg_plano, values=["10 pt", "11 pt", "12 pt"], state="readonly", width=10)
        self.combo_fonte_plano.current(1)
        self.combo_fonte_plano.grid(row=0, column=5, padx=5, pady=5, sticky="w")
        ttk.Label(frame_cfg_plano, text="Entrelinhas:").grid(row=1, column=0, padx=5, pady=5, sticky="e")
        self.combo_entrelinhas_plano = ttk.Combobox(frame_cfg_plano, values=["Compacto (1.0)", "Padrão (1.15)", "Expandido (1.5)"], state="readonly", width=15)
        self.combo_entrelinhas_plano.current(1)
        self.combo_entrelinhas_plano.grid(row=1, column=1, padx=5, pady=5, sticky="w")
        ttk.Label(frame_cfg_plano, text="Espaço Parág.:").grid(row=1, column=2, padx=5, pady=5, sticky="e")
        self.combo_espaco_plano = ttk.Combobox(frame_cfg_plano, values=["Mínimo (4 pt)", "Médio (6 pt)", "Amplo (10 pt)"], state="readonly", width=20)
        self.combo_espaco_plano.current(1)
        self.combo_espaco_plano.grid(row=1, column=3, padx=5, pady=5, sticky="w")

    def setup_tab2(self):
        frame = ttk.LabelFrame(self.tab2, text=" Histórico Pedagógico / Planos Anteriores ", padding=15)
        frame.pack(fill="both", expand=True, padx=15, pady=15)
        ttk.Label(frame, text="Carregue planos de aula das semanas passadas para manter o estilo didático e a sequência cronológica:").pack(anchor="w", pady=(0, 5))
        file_box = ttk.Frame(frame)
        file_box.pack(fill="x", pady=5)
        ent = ttk.Entry(file_box, textvariable=self.file_historico_path, width=60)
        ent.pack(side="left", fill="x", expand=True, padx=(0, 5))
        add_context_menu(ent)
        ttk.Button(file_box, text="Procurar...", command=lambda: self.pick_file(self.file_historico_path, [("Word / Text", "*.docx *.txt")])).pack(side="right")
        ttk.Label(frame, text="Resumo de conteúdos ministrados recentemente:").pack(anchor="w", pady=(15, 5))
        self.txt_historico = scrolledtext.ScrolledText(frame, height=12, font=("Consolas", 10))
        self.txt_historico.pack(fill="both", expand=True)
        add_context_menu(self.txt_historico)
        self.create_field_toolbar(frame, self.txt_historico)

    def setup_tab3(self):
        frame = ttk.LabelFrame(self.tab3, text=" Ingestão do Material da Semana (Livro / Apostila) ", padding=15)
        frame.pack(fill="both", expand=True, padx=15, pady=15)
        ttk.Label(frame, text="Selecione o arquivo em PDF ou DOCX contendo as páginas do livro didático da semana:").pack(anchor="w", pady=(0, 5))
        file_box = ttk.Frame(frame)
        file_box.pack(fill="x", pady=5)
        ent = ttk.Entry(file_box, textvariable=self.file_livro_path, width=60)
        ent.pack(side="left", fill="x", expand=True, padx=(0, 5))
        add_context_menu(ent)
        ttk.Button(file_box, text="Procurar PDF/DOCX...", command=lambda: self.pick_file(self.file_livro_path, [("Documentos", "*.pdf *.docx *.txt")])).pack(side="right")
        ttk.Label(frame, text="Anotações / Trechos copiados do livro (Opcional):").pack(anchor="w", pady=(15, 5))
        self.txt_livro = scrolledtext.ScrolledText(frame, height=12, font=("Consolas", 10))
        self.txt_livro.pack(fill="both", expand=True)
        add_context_menu(self.txt_livro)
        self.create_field_toolbar(frame, self.txt_livro)

    def setup_tab4(self):
        frame = ttk.LabelFrame(self.tab4, text=" Conteúdos Complementares & Diretrizes Didáticas ", padding=15)
        frame.pack(fill="both", expand=True, padx=15, pady=15)
        ttk.Label(frame, text="Gênero Textual da Semana (Ex: Tirinha, Fábula, Poema, Conto, Artigo de Opinião):", font=("Segoe UI", 9, "bold")).pack(anchor="w", pady=(0, 2))
        self.ent_genero = ttk.Entry(frame, width=50)
        self.ent_genero.pack(anchor="w", pady=(0, 10))
        add_context_menu(self.ent_genero)
        ttk.Label(frame, text="Conteúdos Extra-Livro / Temas Complementares da Semana:", font=("Segoe UI", 9, "bold")).pack(anchor="w", pady=(0, 2))
        self.ent_extra = ttk.Entry(frame, width=70)
        self.ent_extra.pack(anchor="w", pady=(0, 10))
        add_context_menu(self.ent_extra)
        ttk.Label(frame, text="Frase Pedagógica Pronta (Ênfase Metodológica Selecionável):", font=("Segoe UI", 9, "bold")).pack(anchor="w", pady=(0, 2))
        self.frases_list = [
            "Focar em Atividades Práticas, Projetos e Aprendizagem Mão na Massa (Maker).",
            "Priorizar Exposição Teórica Interativa + Exercícios de Fixação e Sistematização.",
            "Enfatizar Leitura Crítica, Interpretação Textual e Rodas de Conversa/Debates.",
            "Aplicações Gamificadas, Jogos Pedagógicos em Duplas e Dinâmicas de Grupo.",
            "Foco em Consolidação de Aprendizagem, Diagnóstico e Acompanhamento Individualizado."
        ]
        self.combo_frases = ttk.Combobox(frame, values=self.frases_list, width=75, state="readonly")
        self.combo_frases.current(0)
        self.combo_frases.pack(anchor="w", pady=(0, 15))
        ttk.Label(frame, text="Observações e Orientações Específicas para esta Semana:").pack(anchor="w", pady=(0, 2))
        self.txt_observacoes = scrolledtext.ScrolledText(frame, height=6, font=("Segoe UI", 10))
        self.txt_observacoes.pack(fill="both", expand=True)
        add_context_menu(self.txt_observacoes)
        self.create_field_toolbar(frame, self.txt_observacoes)

    def setup_tab5(self):
        frame = ttk.LabelFrame(self.tab5, text=" Filtros da BNCC e Adequação Cognitiva (EF1, EF2 e EM) ", padding=15)
        frame.pack(fill="both", expand=True, padx=15, pady=15)
        box1 = ttk.Frame(frame)
        box1.pack(fill="x", pady=5)
        ttk.Label(box1, text="Série / Ano Escolar:", font=("Segoe UI", 10, "bold")).pack(side="left", padx=(0, 10))
        self.combo_ano = ttk.Combobox(box1, values=[
            "1º Ano", "2º Ano", "3º Ano", "4º Ano", "5º Ano", "6º Ano", "7º Ano", "8º Ano", "9º Ano", "1º Ano EM", "2º Ano EM", "3º Ano EM"
        ], state="readonly", width=14)
        self.combo_ano.current(4)
        self.combo_ano.pack(side="left", padx=(0, 20))
        ttk.Label(box1, text="Tempo / Duração:", font=("Segoe UI", 10, "bold")).pack(side="left", padx=(0, 10))
        self.combo_duracao = ttk.Combobox(box1, values=[
            "1 Aula (50 minutos)", "Bloco Duplo (100 minutos)", "Semanal Completo (5 Dias de Aula)", "Período Integral"
        ], state="readonly", width=25)
        self.combo_duracao.current(2)
        self.combo_duracao.pack(side="left")
        box2 = ttk.Frame(frame)
        box2.pack(fill="x", pady=10)
        self.var_multi = tk.BooleanVar(value=True)
        ttk.Checkbutton(box2, text="Plano Semanal Multidisciplinar / Integrado", variable=self.var_multi).pack(anchor="w")
        self.var_auto_bncc = tk.BooleanVar(value=True)
        ttk.Checkbutton(box2, text="Filtro Inteligente de Habilidades (Leitura automática do bncc_data.json)", variable=self.var_auto_bncc).pack(anchor="w", pady=5)
        ttk.Label(frame, text="Inserir Códigos Específicos da BNCC Manuais (Opcional - Ex: EF05LP01, EM13LGG101):").pack(anchor="w", pady=(10, 2))
        self.ent_bncc_manual = ttk.Entry(frame, width=60)
        self.ent_bncc_manual.pack(anchor="w")
        add_context_menu(self.ent_bncc_manual)
        info_box = tk.Text(frame, height=6, bg="#F1F5F9", relief="flat", font=("Segoe UI", 9))
        info_box.insert("1.0", "🔒 Suporte Expandido BNCC (Educação Infantil até Ensino Médio):\nO Prof. Aula buscará automaticamente o arquivo 'bncc_data.json' embutido no sistema.")
        info_box.configure(state="disabled")
        info_box.pack(fill="x", pady=(15, 0))

    def setup_tab6(self):
        frame = ttk.LabelFrame(self.tab6, text=" Realidade Local, Perfil da Turma, Recursos & Acessibilidade ", padding=15)
        frame.pack(fill="both", expand=True, padx=15, pady=15)
        ttk.Label(frame, text="Nome da Escola / Identificação da Turma (Opcional):", font=("Segoe UI", 9, "bold")).pack(anchor="w", pady=(0, 2))
        self.ent_escola_nome = ttk.Entry(frame, width=60)
        self.ent_escola_nome.pack(anchor="w", pady=(0, 10))
        add_context_menu(self.ent_escola_nome)
        ttk.Label(frame, text="Perfil da Turma e Inclusão (Ex: Alunos PCD, TDAH, Autismo, Ritmo):", font=("Segoe UI", 9, "bold")).pack(anchor="w", pady=(0, 2))
        self.txt_perfil_turma = scrolledtext.ScrolledText(frame, height=3, font=("Segoe UI", 10))
        self.txt_perfil_turma.pack(fill="x", pady=(0, 10))
        add_context_menu(self.txt_perfil_turma)
        ttk.Label(frame, text="Atalhos Rápidos de Infraestrutura & Restrições de Recursos:", font=("Segoe UI", 9, "bold")).pack(anchor="w", pady=(5, 2))
        preset_box = ttk.Frame(frame)
        preset_box.pack(fill="x", pady=(0, 10))
        self.var_res_lousa_caderno = tk.BooleanVar(value=True)
        self.var_res_sem_impressao = tk.BooleanVar(value=True)
        self.var_res_datashow = tk.BooleanVar(value=False)
        self.var_res_patio = tk.BooleanVar(value=False)
        self.var_res_laboratorio = tk.BooleanVar(value=False)
        col1 = ttk.Frame(preset_box)
        col1.pack(side="left", fill="y", padx=(0, 20))
        ttk.Checkbutton(col1, text="📌 Apenas Quadro / Giz e Caderno", variable=self.var_res_lousa_caderno).pack(anchor="w")
        ttk.Checkbutton(col1, text="🚫 Sem impressões para alunos", variable=self.var_res_sem_impressao).pack(anchor="w")
        col2 = ttk.Frame(preset_box)
        col2.pack(side="left", fill="y")
        ttk.Checkbutton(col2, text="🖥️ Projetor / Datashow Disponível", variable=self.var_res_datashow).pack(anchor="w")
        ttk.Checkbutton(col2, text="🏀 Pátio / Quadra Liberada", variable=self.var_res_patio).pack(anchor="w")
        ttk.Checkbutton(col2, text="💻 Laboratório / Tablets", variable=self.var_res_laboratorio).pack(anchor="w")
        ttk.Label(frame, text="Outros Recursos Específicos ou Materiais (Caixa Adaptável):", font=("Segoe UI", 9, "bold")).pack(anchor="w", pady=(5, 2))
        self.txt_infraestrutura = scrolledtext.ScrolledText(frame, height=2, font=("Segoe UI", 10))
        self.txt_infraestrutura.pack(fill="x", pady=(0, 10))
        add_context_menu(self.txt_infraestrutura)
        ttk.Label(frame, text="Realidade Socioespacial da Escola:", font=("Segoe UI", 9, "bold")).pack(anchor="w", pady=(0, 2))
        self.txt_realidade_local = scrolledtext.ScrolledText(frame, height=3, font=("Segoe UI", 10))
        self.txt_realidade_local.pack(fill="both", expand=True)

    def setup_tab7(self):
        frame = ttk.LabelFrame(self.tab7, text=" Gerador de Atividades Impressas (Folha do Aluno) ", padding=15)
        frame.pack(fill="both", expand=True, padx=15, pady=15)

        ttk.Label(frame, text="1. Selecione o Plano de Aula (.docx) base (Opcional se preencher as Instruções abaixo):").pack(anchor="w", pady=(0, 5))
        file_box = ttk.Frame(frame)
        file_box.pack(fill="x", pady=5)
        ent = ttk.Entry(file_box, textvariable=self.file_plano_base_path, width=60)
        ent.pack(side="left", fill="x", expand=True, padx=(0, 5))
        add_context_menu(ent)
        ttk.Button(file_box, text="Procurar...", command=lambda: self.pick_file(self.file_plano_base_path, [("Word Document", "*.docx")])).pack(side="right")

        ttk.Label(frame, text="2. Instruções Livres para a Folha (Elas têm PRIORIDADE MÁXIMA sobre os botões):", font=("Segoe UI", 9, "bold")).pack(anchor="w", pady=(10, 5))
        self.txt_instrucoes_atividade = scrolledtext.ScrolledText(frame, height=4, font=("Consolas", 10))
        self.txt_instrucoes_atividade.pack(fill="both", expand=True)
        add_context_menu(self.txt_instrucoes_atividade)
        self.create_field_toolbar(frame, self.txt_instrucoes_atividade)

        # ======== NOVO FILTRO DE IA ========
        frame_filtro = ttk.Frame(frame)
        frame_filtro.pack(fill="x", pady=(5, 5))
        ttk.Label(frame_filtro, text="Módulo de Sanitização (Filtro de IA):", font=("Segoe UI", 9, "bold")).pack(side="left", padx=(0, 10))
        self.combo_filtro_modo = ttk.Combobox(frame_filtro, values=["Automático (Recomendado)", "Prévia e Confirmação", "Modo Direto (Sem Filtro)"], state="readonly", width=30)
        self.combo_filtro_modo.current(0)
        self.combo_filtro_modo.pack(side="left")

        # ======== OPÇÕES PEDAGÓGICAS ========
        frame_flags = ttk.LabelFrame(frame, text=" ⚙️ Opções Pedagógicas (Botões de Atalho Padrão) ", padding=10)
        frame_flags.pack(fill="x", pady=(5, 5))

        self.var_gabarito = tk.BooleanVar(value=False)
        ttk.Checkbutton(frame_flags, text="Gerar Gabarito", variable=self.var_gabarito).grid(row=0, column=0, padx=5, pady=5, sticky="w")

        self.var_simplificar = tk.BooleanVar(value=False)
        ttk.Checkbutton(frame_flags, text="Simplificar (AEE)", variable=self.var_simplificar).grid(row=0, column=1, padx=5, pady=5, sticky="w")

        self.var_texto_apoio = tk.BooleanVar(value=False)
        ttk.Checkbutton(frame_flags, text="Gerar Texto Base", variable=self.var_texto_apoio).grid(row=0, column=2, padx=5, pady=5, sticky="w")

        self.combo_tamanho_texto = ttk.Combobox(frame_flags, values=["Automático", "Pequeno", "Médio", "Grande"], state="readonly", width=12)
        self.combo_tamanho_texto.current(0)
        self.combo_tamanho_texto.grid(row=0, column=3, padx=5, pady=5, sticky="w")

        ttk.Label(frame_flags, text="Nível Cognitivo:").grid(row=1, column=0, padx=5, pady=5, sticky="e")
        self.combo_nivel = ttk.Combobox(frame_flags, values=["Automático", "Literal (Fácil)", "Inferencial (Médio)", "Crítico (Avançado)"], state="readonly", width=18)
        self.combo_nivel.current(0)
        self.combo_nivel.grid(row=1, column=1, padx=5, pady=5, sticky="w")

        ttk.Label(frame_flags, text="Foco:").grid(row=1, column=2, padx=5, pady=5, sticky="e")
        self.combo_foco = ttk.Combobox(frame_flags, values=["Automático", "Fixação", "Revisão", "Diagnóstica"], state="readonly", width=15)
        self.combo_foco.current(0)
        self.combo_foco.grid(row=1, column=3, padx=5, pady=5, sticky="w")

        ttk.Label(frame_flags, text="Formato de Resposta:").grid(row=2, column=0, padx=5, pady=5, sticky="e")
        self.combo_formato = ttk.Combobox(frame_flags, values=["Linhas (Padrão)", "Caixa de Rascunho/Cálculo"], state="readonly", width=22)
        self.combo_formato.current(0)
        self.combo_formato.grid(row=2, column=1, columnspan=3, padx=5, pady=5, sticky="w")

        # ======== LAYOUT ========
        frame_cfg = ttk.LabelFrame(frame, text=" Configurações de Impressão e Layout ", padding=10)
        frame_cfg.pack(fill="x", pady=(5, 10))

        ttk.Label(frame_cfg, text="Margens:").grid(row=0, column=0, padx=5, pady=2, sticky="e")
        self.combo_margem = ttk.Combobox(frame_cfg, values=["Normal", "Estreita", "Moderada", "Larga"], state="readonly", width=15)
        self.combo_margem.current(0)
        self.combo_margem.grid(row=0, column=1, padx=5, pady=2, sticky="w")
        ttk.Label(frame_cfg, text="Colunas:").grid(row=0, column=2, padx=5, pady=2, sticky="e")
        self.combo_coluna = ttk.Combobox(frame_cfg, values=["1 Coluna (Padrão)", "2 Colunas (Apostila)"], state="readonly", width=20)
        self.combo_coluna.current(0)
        self.combo_coluna.grid(row=0, column=3, padx=5, pady=2, sticky="w")
        ttk.Label(frame_cfg, text="Fonte:").grid(row=0, column=4, padx=5, pady=2, sticky="e")
        self.combo_fonte = ttk.Combobox(frame_cfg, values=["10 pt", "11 pt", "12 pt"], state="readonly", width=10)
        self.combo_fonte.current(1)
        self.combo_fonte.grid(row=0, column=5, padx=5, pady=2, sticky="w")

        ttk.Label(frame_cfg, text="Entrelinhas:").grid(row=1, column=0, padx=5, pady=2, sticky="e")
        self.combo_entrelinhas = ttk.Combobox(frame_cfg, values=["Compacto (1.0)", "Padrão (1.15)", "Expandido (1.5)"], state="readonly", width=15)
        self.combo_entrelinhas.current(1)
        self.combo_entrelinhas.grid(row=1, column=1, padx=5, pady=2, sticky="w")
        ttk.Label(frame_cfg, text="Espaço Parág.:").grid(row=1, column=2, padx=5, pady=2, sticky="e")
        self.combo_espaco = ttk.Combobox(frame_cfg, values=["Mínimo (4 pt)", "Médio (6 pt)", "Amplo (10 pt)"], state="readonly", width=20)
        self.combo_espaco.current(1)
        self.combo_espaco.grid(row=1, column=3, padx=5, pady=2, sticky="w")

        self.btn_gerar_atividade = tk.Button(
            frame, text="🖨️ GERAR CADERNO DE ATIVIDADES",
            command=self.start_activity_generation_thread,
            bg="#7C3AED", fg="white", font=("Segoe UI", 10, "bold"), padx=12, pady=5, relief="raised"
        )
        self.btn_gerar_atividade.pack(pady=5)

    def create_footer(self):
        footer_frame = tk.Frame(self.root)
        footer_frame.pack(fill="x", padx=10, pady=10)

        self.btn_gerar = tk.Button(footer_frame, text="🚀 GERAR PLANO NO WORD (.DOCX)", command=self.start_generation_thread, bg="#2563EB", fg="white", font=("Segoe UI", 10, "bold"), padx=12, pady=8, relief="raised")
        self.btn_gerar.pack(side="right", padx=4)

        self.btn_revisar = tk.Button(footer_frame, text="✏️ Revisão Cirúrgica", command=self.open_surgical_feedback_dialog, bg="#475569", fg="white", font=("Segoe UI", 10, "bold"), padx=10, pady=8, relief="raised")
        self.btn_revisar.pack(side="right", padx=4)

        # Novo Sistema de Gavetas
        self.btn_salvar_como = tk.Button(footer_frame, text="📂 Salvar Como...", command=self.open_save_as_dialog, bg="#059669", fg="white", font=("Segoe UI", 10, "bold"), padx=10, pady=8, relief="raised")
        self.btn_salvar_como.pack(side="right", padx=4)

        self.btn_salvar_rascunho = tk.Button(footer_frame, text="💾 Salvar", command=self.quick_save_draft, bg="#10B981", fg="white", font=("Segoe UI", 10, "bold"), padx=10, pady=8, relief="raised")
        self.btn_salvar_rascunho.pack(side="right", padx=4)

        self.btn_gavetas = tk.Button(footer_frame, text="📚 Meus Rascunhos", command=self.open_draft_manager, bg="#D97706", fg="white", font=("Segoe UI", 10, "bold"), padx=10, pady=8, relief="raised")
        self.btn_gavetas.pack(side="right", padx=4)

        self.lbl_status = ttk.Label(footer_frame, text="Pronto para trabalhar. Motor Multi-IA Ativo.", font=("Segoe UI", 10))
        self.lbl_status.pack(side="left", padx=10)

    # ================= LOGICA DO PLANO DE AULA =================
    def start_generation_thread(self):
        g_key = self.api_key_var.get().strip()
        gr_key = self.groq_key_var.get().strip()
        op_key = self.openrouter_key_var.get().strip()

        if not g_key and not gr_key and not op_key:
            messagebox.showerror("Erro", "Configure ao menos uma chave de API em Configurações.")
            self.open_settings_dialog()
            return

        cfg_margem_plano = self.combo_margem_plano.get()
        cfg_coluna_plano_str = self.combo_coluna_plano.get()
        cfg_num_colunas_plano = 2 if "2" in cfg_coluna_plano_str else 1
        cfg_fonte_plano_str = self.combo_fonte_plano.get()
        cfg_tamanho_fonte_plano = int(cfg_fonte_plano_str.split()[0])
        mapa_linhas = {"Compacto (1.0)": 1.0, "Padrão (1.15)": 1.15, "Expandido (1.5)": 1.5}
        cfg_entrelinhas_plano = mapa_linhas.get(self.combo_entrelinhas_plano.get(), 1.15)
        mapa_espaco = {"Mínimo (4 pt)": 4, "Médio (6 pt)": 6, "Amplo (10 pt)": 10}
        cfg_espaco_plano = mapa_espaco.get(self.combo_espaco_plano.get(), 6)

        self.btn_gerar.config(state="disabled", bg="#94A3B8")
        self.lbl_status.config(text="⏳ Consultando motor de IA (Gerando JSON)...")

        threading.Thread(target=self.run_generation_process, args=(cfg_margem_plano, cfg_num_colunas_plano, cfg_tamanho_fonte_plano, cfg_entrelinhas_plano, cfg_espaco_plano), daemon=True).start()

    def run_generation_process(self, tipo_margem, num_colunas, tamanho_fonte, entrelinhas, espaco_paragrafo):
        try:
            txt_esqueleto = self.extract_text_from_file(self.file_esqueleto_path.get()) + "\n" + self.txt_esqueleto.get("1.0", tk.END).strip()
            txt_historico = self.extract_text_from_file(self.file_historico_path.get()) + "\n" + self.txt_historico.get("1.0", tk.END).strip()
            txt_livro = self.extract_text_from_file(self.file_livro_path.get()) + "\n" + self.txt_livro.get("1.0", tk.END).strip()

            genero = self.ent_genero.get().strip()
            extra = self.ent_extra.get().strip()
            frase_diretriz = self.combo_frases.get()
            obs = self.txt_observacoes.get("1.0", tk.END).strip()

            ano_selecionado = self.combo_ano.get()
            duracao_selecionada = self.combo_duracao.get()
            is_multi = self.var_multi.get()
            bncc_manual = self.ent_bncc_manual.get().strip()
            escola_nome = self.ent_escola_nome.get().strip()
            perfil_turma = self.txt_perfil_turma.get("1.0", tk.END).strip()
            infraestrutura_texto = self.txt_infraestrutura.get("1.0", tk.END).strip()

            infra_presets = []
            if self.var_res_lousa_caderno.get(): infra_presets.append("APENAS LOUSA/GIZ E CADERNO.")
            if self.var_res_sem_impressao.get(): infra_presets.append("SEM IMPRESSÕES INDIVIDUAIS PARA ALUNOS.")
            if self.var_res_datashow.get(): infra_presets.append("Datashow disponível.")
            if self.var_res_patio.get(): infra_presets.append("Pátio liberado.")
            if self.var_res_laboratorio.get(): infra_presets.append("Laboratório disponível.")
            infra_final = "; ".join(infra_presets)
            if infraestrutura_texto: infra_final += f" | Outros: {infraestrutura_texto}"

            bncc_formatted_text = ""
            if os.path.exists(bncc_json_path) and self.var_auto_bncc.get():
                try:
                    with open(bncc_json_path, 'r', encoding='utf-8') as f: bncc_data = json.load(f)
                    prefixos_bncc = {
                        "1º Ano": ["EF01", "EF12", "EF15"], "2º Ano": ["EF02", "EF12", "EF15"],
                        "3º Ano": ["EF03", "EF35", "EF15"], "4º Ano": ["EF04", "EF35", "EF15"],
                        "5º Ano": ["EF05", "EF35", "EF15"], "6º Ano": ["EF06", "EF67", "EF69"],
                        "7º Ano": ["EF07", "EF67", "EF69"], "8º Ano": ["EF08", "EF89", "EF69"],
                        "9º Ano": ["EF09", "EF89", "EF69"], "1º Ano EM": ["EM13", "EM"],
                        "2º Ano EM": ["EM13", "EM"], "3º Ano EM": ["EM13", "EM"]
                    }
                    prefixos_validos = prefixos_bncc.get(ano_selecionado, [])
                    habilidades_filtradas = [h for h in bncc_data if any(h.get("codigo", "").startswith(p) for p in prefixos_validos)]
                    if habilidades_filtradas:
                        linhas_bncc = [f"[{h.get('codigo', '')}] {h.get('descricao', '').replace(chr(10), ' ')}" for h in habilidades_filtradas]
                        bncc_formatted_text = "\n".join(linhas_bncc)
                except Exception as e: print(f"Erro ao processar JSON: {e}")

            system_instruction = (
                "Você é o motor pedagógico especialista do Prof. Aula.\n"
                "IMPORTANTE: Você DEVE retornar sua resposta EXCLUSIVAMENTE em formato JSON válido, sem texto antes ou depois.\n"
                "{\n"
                "  \"cabecalho\": {\"tema_central\": \"...\", \"disciplinas\": \"...\"},\n"
                "  \"bncc_competencias\": [\"codigo 1: descrição\"],\n"
                "  \"objetivos_aprendizagem\": [\"objetivo 1\"],\n"
                "  \"conteudo_programatico\": [\"conteudo 1\"],\n"
                "  \"metodologia_desenvolvimento\": [{\"etapa\": \"Introdução\", \"descricao\": \"...\"}],\n"
                "  \"recursos_didaticos\": [\"recurso 1\"],\n"
                "  \"avaliacao\": \"...\"\n"
                "}\n"
            )

            prompt_user = f"""
- Série: {ano_selecionado} | Duração: {duracao_selecionada} | Multidisciplinar: {is_multi}
- Escola: {escola_nome} | Perfil: {perfil_turma} | Infra: {infra_final}
- Esqueleto: {txt_esqueleto}
- Histórico: {txt_historico}
- Livro: {txt_livro}
- Gênero: {genero} | Extra: {extra} | Ênfase: {frase_diretriz} | Obs: {obs}
- Habilidades BNCC:\n{bncc_formatted_text}
- Manuais: {bncc_manual}
Gere o plano completo.
"""
            resposta_json_text, telemetria = call_ai_multi_provider(
                gemini_key=self.api_key_var.get().strip(),
                groq_key=self.groq_key_var.get().strip(),
                openrouter_key=self.openrouter_key_var.get().strip(),
                provider_priority=self.provider_priority_var.get().strip(),
                prompt=prompt_user,
                system_instruction=system_instruction,
                temperature=0.6
            )

            plano_dados = self.parse_ai_json(resposta_json_text)
            target_dir = self.save_dir_var.get().strip()
            if not os.path.isdir(target_dir): target_dir = self.load_save_dir()
            timestamp = datetime.datetime.now().strftime("%Y-%m-%d_%H-%M-%S")
            output_filename = os.path.join(target_dir, f"Plano_de_Aula_{ano_selecionado.replace(' ', '_')}_{timestamp}.docx")

            self.export_json_to_docx(plano_dados, output_filename, tipo_margem, num_colunas, tamanho_fonte, entrelinhas, espaco_paragrafo)
            self.file_plano_base_path.set(output_filename)

            self.lbl_status.config(text=f"✅ Sucesso! [{telemetria['model']}]")
            messagebox.showinfo("Sucesso", f"Plano gerado com sucesso!\n\nModelo: {telemetria['model']}\nTokens: {telemetria['total_tokens']}\nSalvo em:\n{output_filename}")

            if os.name == 'posix': os.system(f"xdg-open '{output_filename}'")
            else: os.startfile(output_filename)

        except Exception as e:
            self.lbl_status.config(text="❌ Erro na geração.")
            self.root.after(0, lambda: show_error_dialog(self.root, "Erro na Geração de Plano", str(e), getattr(self, 'is_dark_mode', False)))
        finally:
            self.btn_gerar.config(state="normal", bg="#2563EB")

    # ================= LÓGICA DE ATIVIDADES E FILTRO IA =================
    def start_activity_generation_thread(self):
        g_key = self.api_key_var.get().strip()
        gr_key = self.groq_key_var.get().strip()
        op_key = self.openrouter_key_var.get().strip()

        if not g_key and not gr_key and not op_key:
            messagebox.showerror("Erro", "Configure ao menos uma chave de API em Configurações.")
            return

        plano_path = self.file_plano_base_path.get().strip()
        instrucoes_brutas = self.txt_instrucoes_atividade.get("1.0", tk.END).strip()

        if (not plano_path or not os.path.exists(plano_path)) and not instrucoes_brutas:
            messagebox.showwarning("Aviso", "Forneça um Plano de Aula (.docx) OU digite Instruções na Aba 7.")
            return

        # Captura layouts e flags antes da thread
        layout_cfg = {
            "margem": self.combo_margem.get(),
            "num_colunas": 2 if "2" in self.combo_coluna.get() else 1,
            "tamanho_fonte": int(self.combo_fonte.get().split()[0]),
            "entrelinhas": {"Compacto (1.0)": 1.0, "Padrão (1.15)": 1.15, "Expandido (1.5)": 1.5}.get(self.combo_entrelinhas.get(), 1.15),
            "espaco_paragrafo": {"Mínimo (4 pt)": 4, "Médio (6 pt)": 6, "Amplo (10 pt)": 10}.get(self.combo_espaco.get(), 6),
            "formato_UI": self.combo_formato.get()
        }

        flags_cfg = {
            "opt_gabarito": self.var_gabarito.get(),
            "opt_simplificar": self.var_simplificar.get(),
            "opt_texto_apoio": self.var_texto_apoio.get(),
            "opt_tamanho_texto": self.combo_tamanho_texto.get(),
            "opt_nivel": self.combo_nivel.get(),
            "opt_foco": self.combo_foco.get()
        }

        escola_nome = self.ent_escola_nome.get().strip()
        modo_filtro = self.combo_filtro_modo.get()

        self.btn_gerar_atividade.config(state="disabled", bg="#A78BFA")
        self.lbl_status.config(text="⏳ Analisando instruções e preparando gerador...")

        if modo_filtro == "Modo Direto (Sem Filtro)" or not instrucoes_brutas:
            threading.Thread(target=self.run_activity_main_generation, args=(plano_path, instrucoes_brutas, layout_cfg, flags_cfg, escola_nome), daemon=True).start()
        else:
            # === SISTEMA DE SMART CACHE (Memória Fantasma) ===
            if self.last_raw_instruction and instrucoes_brutas == self.last_raw_instruction and self.cached_clean_instruction:
                print("Usando Cache Fantasma para o filtro de IA.")
                if modo_filtro == "Prévia e Confirmação":
                    self.root.after(0, self.show_preview_dialog, self.cached_clean_data, self.cached_clean_instruction, plano_path, layout_cfg, flags_cfg, escola_nome)
                else:
                    self.root.after(0, lambda: self.lbl_status.config(text="✅ Filtro em Cache recuperado! Gerando documento..."))
                    threading.Thread(target=self.run_activity_main_generation, args=(plano_path, self.cached_clean_instruction, layout_cfg, flags_cfg, escola_nome), daemon=True).start()
            else:
                threading.Thread(target=self.run_ai_filter, args=(instrucoes_brutas, modo_filtro, plano_path, layout_cfg, flags_cfg, escola_nome), daemon=True).start()

    def run_ai_filter(self, instrucoes_brutas, modo_filtro, plano_path, layout_cfg, flags_cfg, escola_nome):
        system_instruction = (
            "Você é o Sanitizador de Prompts do Prof. Aula. Sua função é ler o que o professor ditou/escreveu e converter em um JSON claro.\n"
            "Se o professor pediu uma história/leitura, garanta que isso conste como prioridade.\n"
            "{\n"
            "  \"resumo_formatado\": \"Texto limpo, corrigido e direto do que deve ser feito na atividade (incluindo a ordem de gerar um texto de apoio se pedido).\",\n"
            "  \"texto_leitura\": \"Sim / Não\",\n"
            "  \"qtd_questoes\": \"Estimativa\"\n"
            "}"
        )
        prompt_user = f"Instrução bruta do professor:\n{instrucoes_brutas}"

        max_tentativas = 3
        sucesso = False
        resposta_json = ""
        erro_final = None

        for tentativa in range(1, max_tentativas + 1):
            try:
                self.root.after(0, lambda t=tentativa: self.lbl_status.config(text=f"⏳ Sanitizando instruções (Tentativa {t}/{max_tentativas})..."))

                resposta_json, telemetria = call_ai_multi_provider(
                    gemini_key=self.api_key_var.get().strip(),
                    groq_key=self.groq_key_var.get().strip(),
                    openrouter_key=self.openrouter_key_var.get().strip(),
                    provider_priority=self.provider_priority_var.get().strip(),
                    prompt=prompt_user,
                    system_instruction=system_instruction,
                    temperature=0.3,
                    force_fast_model=True
                )

                sucesso = True
                break

            except Exception as e:
                print(f"Aviso: Filtro falhou na tentativa {tentativa} ({e}).")
                erro_final = e
                if tentativa < max_tentativas:
                    time.sleep(2)

        if not sucesso:
            self.root.after(0, lambda: messagebox.showerror(
                "Erro de Conexão",
                f"Os servidores da IA falharam após {max_tentativas} tentativas de sanitizar o texto.\n\n"
                f"Servidor relatou: {erro_final}\n\n"
                "Aguarde uns instantes e tente novamente, ou mude para o 'Modo Direto' na Aba 7 caso tenha pressa."
            ))
            self.root.after(0, lambda: self.btn_gerar_atividade.config(state="normal", bg="#7C3AED"))
            self.root.after(0, lambda: self.lbl_status.config(text="Geração cancelada. A instrução exigia filtro obrigatório."))
            return

        try:
            dados_filtro = self.parse_ai_json(resposta_json)
            instrucoes_limpas = dados_filtro.get("resumo_formatado", instrucoes_brutas)

            self.last_raw_instruction = instrucoes_brutas
            self.cached_clean_data = dados_filtro
            self.cached_clean_instruction = instrucoes_limpas

            if modo_filtro == "Prévia e Confirmação":
                self.root.after(0, self.show_preview_dialog, dados_filtro, instrucoes_limpas, plano_path, layout_cfg, flags_cfg, escola_nome)
            else:
                self.root.after(0, lambda: self.lbl_status.config(text="✅ Filtro aplicado! Preparando documento..."))
                threading.Thread(target=self.run_activity_main_generation, args=(plano_path, instrucoes_limpas, layout_cfg, flags_cfg, escola_nome), daemon=True).start()

        except Exception as parse_error:
            self.root.after(0, lambda: messagebox.showerror("Erro de Formatação", f"A IA retornou um texto ilegível no filtro.\nErro: {parse_error}"))
            self.root.after(0, lambda: self.btn_gerar_atividade.config(state="normal", bg="#7C3AED"))
            self.root.after(0, lambda: self.lbl_status.config(text="Erro de leitura do Filtro de IA."))

    def show_preview_dialog(self, dados_filtro, instrucoes_limpas, plano_path, layout_cfg, flags_cfg, escola_nome):
        win = tk.Toplevel(self.root)
        win.title("Prévia do Processamento (Filtro de IA)")
        win.geometry("520x480")
        win.minsize(500, 380)
        win.resizable(False, False)

        tk.Label(win, text="O sistema organizou as seguintes instruções:", font=("Segoe UI", 11, "bold")).pack(anchor="w", padx=20, pady=(15, 10))

        frame_info = ttk.Frame(win)
        frame_info.pack(fill="x", padx=20)

        tk.Label(frame_info, text=f"Texto de Leitura (História): {dados_filtro.get('texto_leitura', '?')}", font=("Segoe UI", 9, "bold")).pack(anchor="w", pady=2)
        tk.Label(frame_info, text=f"Volume de Questões: {dados_filtro.get('qtd_questoes', '?')}", font=("Segoe UI", 9, "bold")).pack(anchor="w", pady=2)

        tk.Label(win, text="Resumo da Intenção (Será enviado ao motor final):", font=("Segoe UI", 9, "bold")).pack(anchor="w", padx=20, pady=(15, 2))

        txt_preview = scrolledtext.ScrolledText(win, height=5, font=("Segoe UI", 9))
        txt_preview.insert("1.0", instrucoes_limpas)
        txt_preview.pack(fill="x", padx=20)
        txt_preview.config(state="disabled")

        # ======== MODO DESENVOLVEDOR (RAIO-X CONDICIONAL) ========
        if getattr(self, 'dev_mode_var', tk.BooleanVar(value=False)).get():
            frame_dev = ttk.LabelFrame(win, text=" 🔍 Raio-X do Motor (O que a IA vai ler) ", padding=10)
            frame_dev.pack(fill="x", padx=20, pady=(10, 5))

            texto_status = f"SIM (Tamanho: {flags_cfg['opt_tamanho_texto']})" if flags_cfg['opt_texto_apoio'] else "NÃO (Desativado)"

            dev_text = (
                f"• Gabarito: {'SIM (Ativado)' if flags_cfg['opt_gabarito'] else 'NÃO (Desativado)'}\n"
                f"• Simplificação AEE: {'SIM (Ativado)' if flags_cfg['opt_simplificar'] else 'NÃO (Desativado)'}\n"
                f"• Texto Base/Apoio: {texto_status}\n"
                f"• Nível Cognitivo: {flags_cfg['opt_nivel']}\n"
                f"• Foco da Atividade: {flags_cfg['opt_foco']}"
            )
            tk.Label(frame_dev, text=dev_text, font=("Consolas", 8), justify="left", fg="#475569").pack(anchor="w")
        # ================================================================

        def confirmar():
            win.destroy()
            self.lbl_status.config(text="⏳ Prévia confirmada. Gerando Documento Final...")
            threading.Thread(target=self.run_activity_main_generation, args=(plano_path, instrucoes_limpas, layout_cfg, flags_cfg, escola_nome), daemon=True).start()

        def cancelar():
            self.btn_gerar_atividade.config(state="normal", bg="#7C3AED")
            self.lbl_status.config(text="Processamento cancelado. Nenhuma instrução perdida.")
            win.destroy()

        btn_frame = ttk.Frame(win)
        btn_frame.pack(fill="x", padx=20, pady=15)
        tk.Button(btn_frame, text="✅ Confirmar e Gerar", command=confirmar, bg="#10B981", fg="white", font=("Segoe UI", 10, "bold"), padx=10).pack(side="right")
        tk.Button(btn_frame, text="Voltar", command=cancelar, bg="#64748B", fg="white", font=("Segoe UI", 10, "bold"), padx=10).pack(side="left")

        apply_theme_to_window(win, getattr(self, 'is_dark_mode', False))

    def run_activity_main_generation(self, plano_path, instrucoes_finais, layout_cfg, flags_cfg, escola_nome):
        try:
            self.root.after(0, lambda: self.lbl_status.config(text="⏳ Gerando Atividades via Motor Multi-IA principal..."))

            texto_plano = self.extract_text_from_file(plano_path) if plano_path and os.path.exists(plano_path) else ""
            txt_livro = self.extract_text_from_file(self.file_livro_path.get()) + "\n" + self.txt_livro.get("1.0", tk.END).strip()
            perfil_turma = self.txt_perfil_turma.get("1.0", tk.END).strip()
            ano_selecionado = self.combo_ano.get()

            # Mapeamento do tamanho de texto para uma instrução clara para a IA
            tamanho_map = {
                "Automático": "Automático (o que for mais adequado à faixa etária e ao conteúdo)",
                "Pequeno": "Pequeno (texto curto, bem objetivo, com cerca de 2 a 3 parágrafos)",
                "Médio": "Médio (texto intermediário, com cerca de 3 a 5 parágrafos)",
                "Grande": "Grande (texto longo, profundo e detalhado, com mais de 5 parágrafos)"
            }
            tamanho_instrucao = tamanho_map.get(flags_cfg['opt_tamanho_texto'], "Automático")

            system_instruction = (
                "Você é um Designer Instrucional Pedagógico.\n"
                "IMPORTANTE: Você DEVE retornar sua resposta EXCLUSIVAMENTE em formato JSON válido.\n"
                "{\n"
                "  \"cabecalho_atividade\": {\"escola_nome\": \"...\", \"titulo\": \"...\", \"instrucoes_gerais\": \"...\"},\n"
                "  \"texto_apoio\": \"Texto base completo, dividido em parágrafos (use '\\n' para quebrar os parágrafos). Se não for solicitado, deixe vazio.\",\n"
                "  \"questoes\": [\n"
                "    {\n"
                "      \"numero\": 1,\n"
                "      \"enunciado\": \"...\",\n"
                "      \"tipo\": \"aberta\", // ou 'multipla_escolha', 'verdadeiro_falso'\n"
                "      \"espaco_linhas\": 3,\n"
                "      \"alternativas\": []\n"
                "    }\n"
                "  ],\n"
                "  \"gabarito\": [] // ATENÇÃO: SE A DIRETRIZ FOR 'NÃO' PARA GABARITO, DEIXE ESTA LISTA EXATAMENTE ASSIM: VAZIA.\n"
                "}\n"
            )

            status_gabarito = "SIM (Você deve preencher as respostas esperadas na chave gabarito)" if flags_cfg['opt_gabarito'] else "NÃO (É OBRIGATÓRIO que a chave gabarito retorne como uma lista vazia: [])"
            status_simplificar = "SIM (Use vocabulário super acessível e frases curtas)" if flags_cfg['opt_simplificar'] else "NÃO (Mantenha a linguagem padrão adequada à série)"
            status_texto_apoio = f"SIM (Escreva um texto base/história completo misturando as instruções e os conteúdos do plano/livro. Tamanho exigido: {tamanho_instrucao}. Use múltiplas quebras de linha \\n para parágrafos)" if flags_cfg['opt_texto_apoio'] else "NÃO (Se o professor não pedir explicitamente nas instruções acima, deixe vazio)"

            prompt_user = f"""
- Série: {ano_selecionado} | Escola: {escola_nome} | Perfil: {perfil_turma}
- Plano Base: {texto_plano}
- Material Didático: {txt_livro}

[INSTRUÇÕES DO PROFESSOR - PRIORIDADE MÁXIMA]
{instrucoes_finais if instrucoes_finais else 'Nenhuma.'}

[DIRETRIZES DA INTERFACE (CONFIGURAÇÕES DO SISTEMA)]
ATENÇÃO IA: Siga rigorosamente as chaves de configuração abaixo.
- GERAR GABARITO: {status_gabarito}
- SIMPLIFICAR LINGUAGEM (AEE): {status_simplificar}
- GERAR TEXTO DE APOIO: {status_texto_apoio}
- NÍVEL COGNITIVO: {flags_cfg['opt_nivel']}
- FOCO DA ATIVIDADE: {flags_cfg['opt_foco']}
- FORMATO DE RESPOSTA PREDOMINANTE: {layout_cfg['formato_UI']}

Gere a folha de atividades em JSON. Lembre-se: obedeça as diretrizes de Gabarito e Texto de Apoio acima rigorosamente.
"""
            resposta_json_text, telemetria = call_ai_multi_provider(
                gemini_key=self.api_key_var.get().strip(),
                groq_key=self.groq_key_var.get().strip(),
                openrouter_key=self.openrouter_key_var.get().strip(),
                provider_priority=self.provider_priority_var.get().strip(),
                prompt=prompt_user,
                system_instruction=system_instruction,
                temperature=0.7
            )

            atividade_dados = self.parse_ai_json(resposta_json_text)

            target_dir = self.save_dir_var.get().strip()
            if not os.path.isdir(target_dir): target_dir = self.load_save_dir()
            timestamp = datetime.datetime.now().strftime("%Y-%m-%d_%H-%M-%S")
            output_filename = os.path.join(target_dir, f"Atividades_{ano_selecionado.replace(' ', '_')}_{timestamp}.docx")

            self.export_json_to_docx_atividades(
                atividade_dados, output_filename,
                layout_cfg['margem'], layout_cfg['num_colunas'], layout_cfg['tamanho_fonte'],
                layout_cfg['entrelinhas'], layout_cfg['espaco_paragrafo'], layout_cfg['formato_UI']
            )

            # ==== SUCESSO! HORA DE LIMPAR O SMART CACHE ====
            self.last_raw_instruction = None
            self.cached_clean_data = None
            self.cached_clean_instruction = None
            print("Geração da atividade concluída. Cache de filtro invalidado.")
            # ===============================================

            self.root.after(0, lambda: self.lbl_status.config(text=f"✅ Atividades Geradas! [{telemetria['model']}]"))
            self.root.after(0, lambda: messagebox.showinfo("Sucesso", f"Folha de atividades gerada!\n\nModelo: {telemetria['model']}\nTokens: {telemetria['total_tokens']}\nSalva em:\n{output_filename}"))

            if os.name == 'posix': os.system(f"xdg-open '{output_filename}'")
            else: os.startfile(output_filename)

        except Exception as e:
            self.root.after(0, lambda: self.lbl_status.config(text="❌ Erro. A instrução na aba 7 foi preservada para você tentar novamente."))
            self.root.after(0, lambda: show_error_dialog(self.root, "Erro na Geração de Atividades", str(e), getattr(self, 'is_dark_mode', False)))
        finally:
            self.root.after(0, lambda: self.btn_gerar_atividade.config(state="normal", bg="#7C3AED"))

    # ================= MÓDULO CIRÚRGICO =================
    def open_surgical_feedback_dialog(self):
        target_dir = self.save_dir_var.get().strip()
        if not os.path.isdir(target_dir): target_dir = self.load_save_dir()

        docx_files = [f for f in os.listdir(target_dir) if f.endswith(".docx")]

        win = tk.Toplevel(self.root)
        win.title("Prof. Aula — Módulo de Feedback Cirúrgico")
        win.geometry("650x580")
        win.minsize(600, 500)

        tk.Label(win, text="Selecione o Documento gerado recentemente (.docx):", font=("Segoe UI", 10, "bold")).pack(anchor="w", padx=20, pady=(15, 5))

        file_combo = ttk.Combobox(win, values=docx_files, width=70, state="readonly")
        if docx_files: file_combo.current(0)
        file_combo.pack(anchor="w", padx=20, pady=(0, 10))

        tk.Label(win, text="Trecho ou Seção a ser Ajustada:", font=("Segoe UI", 10, "bold")).pack(anchor="w", padx=20, pady=(5, 5))

        txt_trecho = scrolledtext.ScrolledText(win, height=6, font=("Consolas", 9))
        txt_trecho.pack(fill="x", padx=20, pady=(0, 5))
        add_context_menu(txt_trecho)
        self.create_field_toolbar(win, txt_trecho)

        tk.Label(win, text="Instrução de Ajuste Cirúrgico:", font=("Segoe UI", 10, "bold")).pack(anchor="w", padx=20, pady=(5, 5))

        ent_ajuste = ttk.Entry(win, width=75)
        ent_ajuste.pack(anchor="w", padx=20, pady=(0, 15))
        add_context_menu(ent_ajuste)

        status_lbl = ttk.Label(win, text="Pronto para reescrever trecho.", font=("Segoe UI", 9))
        status_lbl.pack(anchor="w", padx=20, pady=(0, 10))

        def executar_revisao_pontual():
            selected_file = file_combo.get()
            trecho_alvo = txt_trecho.get("1.0", tk.END).strip()
            instrucao = ent_ajuste.get().strip()

            if not selected_file or not instrucao:
                messagebox.showerror("Erro", "Selecione um arquivo e informe a instrução.", parent=win)
                return

            filepath = os.path.join(target_dir, selected_file)
            texto_atual = self.extract_text_from_file(filepath)

            status_lbl.config(text="⏳ Aplicando revisão...")
            win.update()

            try:
                system_instruction = "Você é o editor pedagógico do Prof. Aula. Reescreva APENAS o trecho solicitado mantendo a coerência. Retorne apenas o texto final corrigido, sem markdown."
                prompt_surgical = f"[DOCUMENTO]\n{texto_atual}\n\n[TRECHO ALVO]\n{trecho_alvo}\n\n[AJUSTE]\n{instrucao}"

                novo_texto, telemetria = call_ai_multi_provider(
                    gemini_key=self.api_key_var.get().strip(),
                    groq_key=self.groq_key_var.get().strip(),
                    openrouter_key=self.openrouter_key_var.get().strip(),
                    provider_priority=self.provider_priority_var.get().strip(),
                    prompt=prompt_surgical,
                    system_instruction=system_instruction,
                    temperature=0.5
                )

                base_name, ext = os.path.splitext(filepath)
                timestamp = datetime.datetime.now().strftime("%Y-%m-%d_%H-%M-%S")
                novo_filepath = f"{base_name}_revisado_{timestamp}{ext}"

                texto_final = texto_atual.replace(trecho_alvo, novo_texto)

                doc = docx.Document()
                for section in doc.sections:
                    section.top_margin = Inches(0.8)
                    section.bottom_margin = Inches(0.8)
                    section.left_margin = Inches(0.8)
                    section.right_margin = Inches(0.8)

                for linha in texto_final.split('\n'):
                     doc.add_paragraph(linha)

                doc.save(novo_filepath)

                status_lbl.config(text=f"✅ OK! [{telemetria['model']}]")
                messagebox.showinfo("Sucesso", f"Documento revisado salvo em:\n{novo_filepath}", parent=win)

                if os.name == 'posix': os.system(f"xdg-open '{novo_filepath}'")
                else: os.startfile(novo_filepath)
                win.destroy()

            except Exception as e:
                status_lbl.config(text="❌ Erro na revisão.")
                show_error_dialog(win, "Erro na Revisão Cirúrgica", str(e), getattr(self, 'is_dark_mode', False))

        tk.Button(win, text="⚡ Executar Revisão Cirúrgica", command=executar_revisao_pontual, bg="#2563EB", fg="white", font=("Segoe UI", 10, "bold"), padx=15, pady=5).pack(pady=5)
        apply_theme_to_window(win, self.is_dark_mode)

    # ================= EXPORTAÇÃO DOCX =================
    def format_heading(self, doc, text, level=1):
        p = doc.add_heading(level=level)
        run = p.add_run(text)
        run.font.name = "Arial"
        if level == 0:
            run.font.size = Pt(14)
            run.font.color.rgb = RGBColor(0x1E, 0x3A, 0x8A)
        elif level == 1:
            run.font.size = Pt(12)
            run.font.color.rgb = RGBColor(0x25, 0x63, 0xEB)
        else:
            run.font.size = Pt(11)
            run.font.color.rgb = RGBColor(0x33, 0x41, 0x55)
        return p

    def export_json_to_docx(self, json_data, filename, tipo_margem="Normal", num_colunas=1, tamanho_fonte=11, entrelinhas=1.15, espaco_paragrafo=6):
        if not docx: raise ImportError("A biblioteca python-docx não está instalada.")
        doc = docx.Document()
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Arial'
        font.size = Pt(tamanho_fonte)

        p_format = style.paragraph_format
        p_format.line_spacing = entrelinhas
        p_format.space_after = Pt(espaco_paragrafo)
        p_format.space_before = Pt(0)

        section = doc.sections[0]
        margens_map = {
            "Normal": {"top": 2.5, "bottom": 2.5, "left": 2.5, "right": 2.5},
            "Estreita": {"top": 1.27, "bottom": 1.27, "left": 1.27, "right": 1.27},
            "Moderada": {"top": 2.54, "bottom": 2.54, "left": 1.91, "right": 1.91},
            "Larga": {"top": 2.54, "bottom": 2.54, "left": 5.08, "right": 5.08}
        }
        cfg = margens_map.get(tipo_margem, margens_map["Normal"])

        section.top_margin = Cm(cfg["top"])
        section.bottom_margin = Cm(cfg["bottom"])
        section.left_margin = Cm(cfg["left"])
        section.right_margin = Cm(cfg["right"])

        if num_colunas == 2:
            sectPr = section._sectPr
            cols = sectPr.xpath("./w:cols")
            cols_elm = cols[0] if cols else OxmlElement("w:cols")
            cols_elm.set(qn("w:num"), "2")
            cols_elm.set(qn("w:space"), "720")
            if not cols: sectPr.append(cols_elm)

        cabecalho = json_data.get("cabecalho", {})
        self.format_heading(doc, "Plano de Aula", level=0)

        p = doc.add_paragraph()
        p.add_run("Tema Central: ").bold = True
        p.add_run(sanitize_text(cabecalho.get("tema_central", "Não definido")))
        p = doc.add_paragraph()
        p.add_run("Disciplinas/Áreas: ").bold = True
        p.add_run(sanitize_text(cabecalho.get("disciplinas", "Não definido")))
        doc.add_paragraph()

        bncc_list = json_data.get("bncc_competencias", [])
        if bncc_list:
            self.format_heading(doc, "Habilidades e Competências (BNCC)", level=1)
            for item in bncc_list: doc.add_paragraph(sanitize_text(item), style='List Bullet')
            doc.add_paragraph()

        objetivos = json_data.get("objetivos_aprendizagem", [])
        if objetivos:
            self.format_heading(doc, "Objetivos de Aprendizagem", level=1)
            for obj in objetivos: doc.add_paragraph(sanitize_text(obj), style='List Bullet')
            doc.add_paragraph()

        conteudo = json_data.get("conteudo_programatico", [])
        if conteudo:
            self.format_heading(doc, "Conteúdo Programático", level=1)
            for cont in conteudo: doc.add_paragraph(sanitize_text(cont), style='List Bullet')
            doc.add_paragraph()

        metodologia = json_data.get("metodologia_desenvolvimento", [])
        if metodologia:
            self.format_heading(doc, "Desenvolvimento Metodológico", level=1)
            for etapa in metodologia:
                p = doc.add_paragraph()
                p.add_run(f"{sanitize_text(etapa.get('etapa', 'Etapa'))}: ").bold = True
                p.add_run(sanitize_text(etapa.get('descricao', '')))
            doc.add_paragraph()

        recursos = json_data.get("recursos_didaticos", [])
        if recursos:
            self.format_heading(doc, "Recursos Didáticos", level=1)
            p = doc.add_paragraph(", ".join([sanitize_text(r) for r in recursos]))
            doc.add_paragraph()

        avaliacao = json_data.get("avaliacao", "")
        if avaliacao:
            self.format_heading(doc, "Avaliação", level=1)
            doc.add_paragraph(sanitize_text(avaliacao))

        doc.save(filename)

    def export_json_to_docx_atividades(self, json_data, filename, tipo_margem="Estreita", num_colunas=2, tamanho_fonte=11, entrelinhas=1.15, espaco_paragrafo=4, formato_UI="Linhas"):
        if not docx: raise ImportError("A biblioteca python-docx não está instalada.")
        doc = docx.Document()
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Arial'
        font.size = Pt(tamanho_fonte)

        p_format = style.paragraph_format
        p_format.line_spacing = entrelinhas
        p_format.space_after = Pt(espaco_paragrafo)
        p_format.space_before = Pt(0)

        section = doc.sections[0]
        margens_map = {
            "Normal": {"top": 2.5, "bottom": 2.5, "left": 2.5, "right": 2.5},
            "Estreita": {"top": 1.27, "bottom": 1.27, "left": 1.27, "right": 1.27},
            "Moderada": {"top": 2.54, "bottom": 2.54, "left": 1.91, "right": 1.91},
            "Larga": {"top": 2.54, "bottom": 2.54, "left": 5.08, "right": 5.08}
        }
        cfg = margens_map.get(tipo_margem, margens_map["Estreita"])

        section.top_margin = Cm(cfg["top"])
        section.bottom_margin = Cm(cfg["bottom"])
        section.left_margin = Cm(cfg["left"])
        section.right_margin = Cm(cfg["right"])

        if num_colunas == 2:
            sectPr = section._sectPr
            cols = sectPr.xpath("./w:cols")
            cols_elm = cols[0] if cols else OxmlElement("w:cols")
            cols_elm.set(qn("w:num"), "2")
            cols_elm.set(qn("w:space"), "720")
            if not cols: sectPr.append(cols_elm)

        largura_util = 21.0 - (cfg["left"] + cfg["right"])
        largura_linha = (largura_util - 1.27) / 2 if num_colunas == 2 else largura_util

        cabecalho = json_data.get("cabecalho_atividade", {})
        escola_txt = cabecalho.get("escola_nome", "").strip()
        if not escola_txt or escola_txt.lower() == "null" or escola_txt == "...":
            escola_txt = "____________________________________"

        p_cab = doc.add_paragraph()
        p_cab.add_run(f"ESCOLA: {escola_txt}\n").bold = True
        p_cab.add_run("NOME:\n").bold = True
        p_cab.add_run("DATA: ____/____/________   TURMA: _______").bold = True
        doc.add_paragraph()
        self.format_heading(doc, sanitize_text(cabecalho.get("titulo", "Atividades")), level=0)

        instrucoes = cabecalho.get("instrucoes_gerais", "")
        if instrucoes:
            p_inst = doc.add_paragraph()
            p_inst.add_run(sanitize_text(instrucoes)).italic = True
            doc.add_paragraph()

        texto_apoio = json_data.get("texto_apoio", "").strip()
        if texto_apoio and texto_apoio.lower() != "null":
            p_apoio_title = doc.add_paragraph()
            p_apoio_title.add_run("Texto Base / Leitura:").bold = True

            # ======== CORREÇÃO DE PARÁGRAFOS DO TEXTO DE APOIO ========
            # Quebra o texto da IA usando o ENTER (\n)
            for paragrafo_texto in texto_apoio.split('\n'):
                linha_limpa = sanitize_text(paragrafo_texto)
                # Se não for linha vazia, cria um novo parágrafo autêntico no Word
                if linha_limpa:
                    p_apoio = doc.add_paragraph()
                    # Adiciona recuo de primeira linha (1.25 cm = "espaço de dois dedos")
                    p_apoio.paragraph_format.first_line_indent = Cm(1.25)
                    # Justifica o texto para ficar com visual de livro
                    p_apoio.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    p_apoio.add_run(linha_limpa)

            doc.add_paragraph() # Dá um espaço antes de iniciar as questões

        for q in json_data.get("questoes", []):
            enunciado_limpo = sanitize_text(q.get("enunciado", ""))
            p_q = doc.add_paragraph()
            p_q.paragraph_format.keep_with_next = True
            p_q.add_run(f"Questão {q.get('numero', '')}: {enunciado_limpo}").bold = True

            tipo = q.get("tipo", "aberta")
            if tipo in ["verdadeiro_falso", "multipla_escolha"]:
                for alt in q.get("alternativas", []):
                    alt_limpa = sanitize_text(alt)
                    p_alt = doc.add_paragraph()
                    if tipo == "multipla_escolha": p_alt.add_run(f"•   (   ) {alt_limpa}")
                    else: p_alt.add_run(f"(   ) {alt_limpa}")

            num_linhas = q.get("espaco_linhas", 0)
            if tipo == "aberta" and num_linhas > 0:
                if "Caixa" in formato_UI:
                    table = doc.add_table(rows=1, cols=1)
                    table.style = 'Table Grid'
                    tr = table.rows[0].cells[0]
                    for _ in range(num_linhas): tr.add_paragraph("")
                    doc.add_paragraph()
                else:
                    for _ in range(num_linhas):
                        p_linha = doc.add_paragraph()
                        p_linha.paragraph_format.tab_stops.add_tab_stop(Cm(largura_linha), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.LINES)
                        p_linha.add_run("\t")

        gabarito = json_data.get("gabarito", [])
        if gabarito:
            doc.add_page_break()
            self.format_heading(doc, "Gabarito e Respostas Esperadas", level=1)
            for resp in gabarito:
                p_gab = doc.add_paragraph()
                p_gab.add_run(f"Questão {resp.get('numero', '')}: ").bold = True
                p_gab.add_run(sanitize_text(resp.get("resposta_esperada", "")))

        doc.save(filename)

if __name__ == "__main__":
    root = tk.Tk()
    app = ProfAulaApp(root)
    root.mainloop()
