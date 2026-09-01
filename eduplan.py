import os
import sys
import re
import json
import threading
import time
import datetime
import urllib.request
import urllib.error
import tkinter as tk
from tkinter import ttk, filedialog, messagebox, scrolledtext

# Importação defensiva de bibliotecas externas
try:
    import docx
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    docx = None

try:
    import pypdf
except ImportError:
    pypdf = None

try:
    from google import genai
    from google.genai import types
except ImportError:
    genai = None
    types = None

# ==============================================================================
# GERENCIAMENTO DE DIRETÓRIOS, ASSETS E CONFIGURAÇÕES SEGURAS
# ==============================================================================
if getattr(sys, 'frozen', False):
    BUNDLE_DIR = sys._MEIPASS
    BASE_DIR = os.path.dirname(os.path.abspath(sys.executable))
else:
    BUNDLE_DIR = os.path.dirname(os.path.abspath(__file__))
    BASE_DIR = os.path.dirname(os.path.abspath(__file__))

# Banco de dados da BNCC embutido no executável
bncc_json_path = os.path.join(BUNDLE_DIR, "bncc_data.json")

# Dados pessoais, chaves de API e rascunhos isolados na pasta oculta do usuário (Universal)
DATA_DIR = os.path.join(os.path.expanduser("~"), ".eduplan_ai")
os.makedirs(DATA_DIR, exist_ok=True)

API_KEY_FILE = os.path.join(DATA_DIR, ".gemini_api_key.txt")
GROQ_KEY_FILE = os.path.join(DATA_DIR, ".groq_api_key.txt")
OPENROUTER_KEY_FILE = os.path.join(DATA_DIR, ".openrouter_api_key.txt")
CONFIG_DIR_FILE = os.path.join(DATA_DIR, ".eduplan_dir.txt")
PROVIDER_CHOICE_FILE = os.path.join(DATA_DIR, ".eduplan_provider.txt")
STATE_FILE = os.path.join(DATA_DIR, ".eduplan_state.json")

# ==============================================================================
# MOTOR MULTI-PROVEDOR E RESILIENTE DE CHAMADA DAS IAs (GEMINI, GROQ, OPENROUTER)
# ==============================================================================
def call_openai_compatible_api(endpoint_url, api_key, model_name, prompt, system_instruction, temperature=0.6):
    headers = {
        "Authorization": f"Bearer {api_key}",
        "Content-Type": "application/json"
    }
    payload = {
        "model": model_name,
        "messages": [
            {"role": "system", "content": system_instruction},
            {"role": "user", "content": prompt}
        ],
        "temperature": temperature
    }

    data_bytes = json.dumps(payload).encode("utf-8")
    req = urllib.request.Request(endpoint_url, data=data_bytes, headers=headers, method="POST")

    with urllib.request.urlopen(req, timeout=60) as resp:
        res_data = json.loads(resp.read().decode("utf-8"))
        text = res_data["choices"][0]["message"]["content"]
        usage = res_data.get("usage", {})
        prompt_tokens = usage.get("prompt_tokens", 0)
        completion_tokens = usage.get("completion_tokens", 0)
        total_tokens = usage.get("total_tokens", prompt_tokens + completion_tokens)

        telemetria = {
            "model": f"{model_name}",
            "prompt_tokens": prompt_tokens,
            "response_tokens": completion_tokens,
            "total_tokens": total_tokens
        }
        return text, telemetria

def call_ai_multi_provider(gemini_key, groq_key, openrouter_key, provider_priority, prompt, system_instruction, temperature=0.6):
    providers_queue = []

    if provider_priority == "Groq (Llama 3)" and groq_key:
        providers_queue.append("groq")
    elif provider_priority == "OpenRouter (Multi-IA)" and openrouter_key:
        providers_queue.append("openrouter")
    elif provider_priority == "Google Gemini" and gemini_key:
        providers_queue.append("gemini")

    if "gemini" not in providers_queue and gemini_key and genai:
        providers_queue.append("gemini")
    if "groq" not in providers_queue and groq_key:
        providers_queue.append("groq")
    if "openrouter" not in providers_queue and openrouter_key:
        providers_queue.append("openrouter")

    if not providers_queue:
        raise RuntimeError("Nenhuma chave de API válida configurada. Configure ao menos uma chave em ⚙️ Configurações.")

    errors_log = []

    for provider in providers_queue:
        if provider == "gemini" and gemini_key and genai:
            client = genai.Client(api_key=gemini_key)
            for model_name in ["gemini-3.6-flash"]:
                for attempt in range(3):
                    try:
                        response = client.models.generate_content(
                            model=model_name,
                            contents=prompt,
                            config=types.GenerateContentConfig(
                                system_instruction=system_instruction,
                                temperature=temperature,
                            )
                        )
                        if response and response.text:
                            usage = getattr(response, "usage_metadata", None)
                            p_tok = getattr(usage, "prompt_token_count", 0) if usage else 0
                            r_tok = getattr(usage, "candidates_token_count", 0) if usage else 0
                            t_tok = getattr(usage, "total_token_count", p_tok + r_tok) if usage else 0
                            return response.text, {"model": f"Gemini ({model_name})", "prompt_tokens": p_tok, "response_tokens": r_tok, "total_tokens": t_tok}
                    except Exception as e:
                        err_msg = str(e).upper()
                        if "503" in err_msg or "UNAVAILABLE" in err_msg:
                            errors_log.append(f"Gemini [{model_name}] (Tentativa {attempt+1}): Sobrecarga/503. Aguardando...")
                            time.sleep(2)
                            continue
                        else:
                            errors_log.append(f"Gemini [{model_name}]: {e}")
                            break

        elif provider == "groq" and groq_key:
            for model_name in ["llama3-8b-8192", "mixtral-8x7b-32768", "llama3-70b-8192"]:
                try:
                    return call_openai_compatible_api(
                        endpoint_url="https://api.groq.com/openai/v1/chat/completions",
                        api_key=groq_key,
                        model_name=model_name,
                        prompt=prompt,
                        system_instruction=system_instruction,
                        temperature=temperature
                    )
                except Exception as e:
                    errors_log.append(f"Groq [{model_name}]: {e}")

        elif provider == "openrouter" and openrouter_key:
            for model_name in ["google/gemma-2-9b-it:free", "meta-llama/llama-3-8b-instruct:free", "huggingfaceh4/zephyr-7b-beta:free"]:
                try:
                    return call_openai_compatible_api(
                        endpoint_url="https://openrouter.ai/api/v1/chat/completions",
                        api_key=openrouter_key,
                        model_name=model_name,
                        prompt=prompt,
                        system_instruction=system_instruction,
                        temperature=temperature
                    )
                except Exception as e:
                    errors_log.append(f"OpenRouter [{model_name}]: {e}")

    raise RuntimeError(f"Todos os provedores configurados falharam no momento.\n\nDetalhes de Diagnóstico:\n" + "\n".join(errors_log))

# ==============================================================================
# HELPER DE INTERFACE: MENU DE CONTEXTO E ATALHOS DE EDIÇÃO DE TEXTO
# ==============================================================================
def add_context_menu(widget):
    menu = tk.Menu(widget, tearoff=0)
    menu.add_command(label="✂️ Cortar", command=lambda: widget.event_generate("<<Cut>>"))
    menu.add_command(label="📋 Copiar", command=lambda: widget.event_generate("<<Copy>>"))
    menu.add_command(label="📥 Colar", command=lambda: widget.event_generate("<<Paste>>"))
    menu.add_separator()
    menu.add_command(label="🔍 Selecionar Tudo", command=lambda: select_all(widget))

    def show_menu(event):
        menu.tk_popup(event.x_root, event.y_root)

    widget.bind("<Button-3>", show_menu)
    widget.bind("<Control-a>", lambda e: select_all(widget))
    widget.bind("<Control-A>", lambda e: select_all(widget))

def select_all(widget):
    if isinstance(widget, (tk.Entry, ttk.Entry)):
        widget.select_range(0, tk.END)
        widget.icursor(tk.END)
    elif isinstance(widget, (tk.Text, scrolledtext.ScrolledText)):
        widget.tag_add(tk.SEL, "1.0", tk.END)
        widget.mark_set(tk.INSERT, "1.0")
        widget.see(tk.INSERT)
    return "break"

def paste_clipboard_to_widget(widget):
    try:
        content = widget.clipboard_get()
        if isinstance(widget, (tk.Entry, ttk.Entry)):
            widget.insert(tk.INSERT, content)
        elif isinstance(widget, (tk.Text, scrolledtext.ScrolledText)):
            widget.insert(tk.INSERT, content)
    except Exception:
        pass

def clear_widget(widget):
    if isinstance(widget, (tk.Entry, ttk.Entry)):
        widget.delete(0, tk.END)
    elif isinstance(widget, (tk.Text, scrolledtext.ScrolledText)):
        widget.delete("1.0", tk.END)


# ==============================================================================
# CLASSE PRINCIPAL DA APLICAÇÃO GUI
# ==============================================================================
class EduPlanAIApp:
    def __init__(self, root):
        self.root = root
        self.root.title("EduPlan AI — Gerador Inteligente de Planos & Atividades (Beta)")
        self.root.geometry("1020x840")
        self.root.minsize(900, 740)

        self.style = ttk.Style()
        self.style.theme_use("clam")

        self.file_esqueleto_path = tk.StringVar()
        self.file_historico_path = tk.StringVar()
        self.file_livro_path = tk.StringVar()
        self.file_plano_base_path = tk.StringVar()

        self.api_key_var = tk.StringVar(value=self.load_key(API_KEY_FILE))
        self.groq_key_var = tk.StringVar(value=self.load_key(GROQ_KEY_FILE))
        self.openrouter_key_var = tk.StringVar(value=self.load_key(OPENROUTER_KEY_FILE))
        self.save_dir_var = tk.StringVar(value=self.load_save_dir())
        self.provider_priority_var = tk.StringVar(value=self.load_provider_choice())

        self.create_header()
        self.create_tabs()
        self.create_footer()

        self.load_state()

    def create_header(self):
        header_frame = tk.Frame(self.root, bg="#1E293B")
        header_frame.pack(fill="x")

        title_lbl = tk.Label(
            header_frame,
            text="EduPlan AI — Automação Pedagógica Multi-IA",
            font=("Segoe UI", 14, "bold"),
            fg="#F8FAFC",
            bg="#1E293B"
        )
        title_lbl.pack(side="left", padx=15, pady=15)

        config_btn = tk.Button(
            header_frame,
            text="⚙️ Configurações Multi-IA",
            command=self.open_settings_dialog,
            bg="#3B82F6",
            fg="white",
            font=("Segoe UI", 9, "bold"),
            relief="flat",
            padx=10
        )
        config_btn.pack(side="right", padx=15, pady=15)

    def load_key(self, filepath):
        if os.path.exists(filepath):
            try:
                with open(filepath, "r", encoding="utf-8") as f:
                    return f.read().strip()
            except Exception:
                return ""
        return ""

    def load_provider_choice(self):
        if os.path.exists(PROVIDER_CHOICE_FILE):
            try:
                with open(PROVIDER_CHOICE_FILE, "r", encoding="utf-8") as f:
                    return f.read().strip()
            except Exception:
                pass
        return "Auto (Fallback Inteligente)"

    def load_save_dir(self):
        if os.path.exists(CONFIG_DIR_FILE):
            try:
                with open(CONFIG_DIR_FILE, "r", encoding="utf-8") as f:
                    path = f.read().strip()
                    if os.path.isdir(path):
                        return path
            except Exception:
                pass
        return BASE_DIR

    def open_settings_dialog(self):
        win = tk.Toplevel(self.root)
        win.title("Configurações Multi-IA do EduPlan AI")
        win.geometry("580x420")
        win.resizable(False, False)

        tk.Label(win, text="Chave Google Gemini API:", font=("Segoe UI", 9, "bold")).pack(anchor="w", padx=20, pady=(12, 2))
        k1 = ttk.Entry(win, textvariable=self.api_key_var, width=65, show="*")
        k1.pack(anchor="w", padx=20, pady=(0, 5))
        add_context_menu(k1)

        tk.Label(win, text="Chave Groq API (Gratuito - Llama 3):", font=("Segoe UI", 9, "bold")).pack(anchor="w", padx=20, pady=(5, 2))
        k2 = ttk.Entry(win, textvariable=self.groq_key_var, width=65, show="*")
        k2.pack(anchor="w", padx=20, pady=(0, 5))
        add_context_menu(k2)

        tk.Label(win, text="Chave OpenRouter API (Gratuito - Multi-Provedores):", font=("Segoe UI", 9, "bold")).pack(anchor="w", padx=20, pady=(5, 2))
        k3 = ttk.Entry(win, textvariable=self.openrouter_key_var, width=65, show="*")
        k3.pack(anchor="w", padx=20, pady=(0, 10))
        add_context_menu(k3)

        tk.Label(win, text="Provedor de IA Prioritário:", font=("Segoe UI", 9, "bold")).pack(anchor="w", padx=20, pady=(2, 2))
        cb_prov = ttk.Combobox(
            win,
            textvariable=self.provider_priority_var,
            values=["Auto (Fallback Inteligente)", "Google Gemini", "Groq (Llama 3)", "OpenRouter (Multi-IA)"],
            state="readonly",
            width=35
        )
        cb_prov.pack(anchor="w", padx=20, pady=(0, 10))

        tk.Label(win, text="Pasta Padrão para Salvar os Planos de Aula (.docx):", font=("Segoe UI", 9, "bold")).pack(anchor="w", padx=20, pady=(5, 2))
        dir_box = ttk.Frame(win)
        dir_box.pack(fill="x", padx=20, pady=(0, 15))

        dir_entry = ttk.Entry(dir_box, textvariable=self.save_dir_var, width=50)
        dir_entry.pack(side="left", fill="x", expand=True, padx=(0, 5))
        add_context_menu(dir_entry)

        def pick_dir():
            chosen_dir = filedialog.askdirectory(initialdir=self.save_dir_var.get())
            if chosen_dir:
                self.save_dir_var.set(chosen_dir)

        ttk.Button(dir_box, text="Procurar...", command=pick_dir).pack(side="right")

        def save_all():
            with open(API_KEY_FILE, "w", encoding="utf-8") as f: f.write(self.api_key_var.get().strip())
            with open(GROQ_KEY_FILE, "w", encoding="utf-8") as f: f.write(self.groq_key_var.get().strip())
            with open(OPENROUTER_KEY_FILE, "w", encoding="utf-8") as f: f.write(self.openrouter_key_var.get().strip())
            with open(PROVIDER_CHOICE_FILE, "w", encoding="utf-8") as f: f.write(self.provider_priority_var.get().strip())

            directory = self.save_dir_var.get().strip()
            if not directory or not os.path.isdir(directory):
                directory = BASE_DIR
                self.save_dir_var.set(directory)

            with open(CONFIG_DIR_FILE, "w", encoding="utf-8") as f: f.write(directory)

            messagebox.showinfo("Sucesso", "Configurações de IA salvas com sucesso!", parent=win)
            win.destroy()

        tk.Button(win, text="Salvar Configurações", command=save_all, bg="#10B981", fg="white", font=("Segoe UI", 10, "bold"), padx=15).pack(pady=5)

    def save_state(self):
        data = {
            "file_esqueleto": self.file_esqueleto_path.get(),
            "esqueleto": self.txt_esqueleto.get("1.0", tk.END).strip(),
            "file_historico": self.file_historico_path.get(),
            "historico": self.txt_historico.get("1.0", tk.END).strip(),
            "file_livro": self.file_livro_path.get(),
            "livro": self.txt_livro.get("1.0", tk.END).strip(),
            "file_plano_base": self.file_plano_base_path.get(),
            "instrucoes_atividade": self.txt_instrucoes_atividade.get("1.0", tk.END).strip(),
            "genero": self.ent_genero.get(),
            "extra": self.ent_extra.get(),
            "frase": self.combo_frases.get(),
            "observacoes": self.txt_observacoes.get("1.0", tk.END).strip(),
            "ano": self.combo_ano.get(),
            "duracao": self.combo_duracao.get(),
            "multi": self.var_multi.get(),
            "auto_bncc": self.var_auto_bncc.get(),
            "bncc_manual": self.ent_bncc_manual.get(),
            "escola_nome": self.ent_escola_nome.get(),
            "perfil_turma": self.txt_perfil_turma.get("1.0", tk.END).strip(),
            "infraestrutura": self.txt_infraestrutura.get("1.0", tk.END).strip(),
            "realidade_local": self.txt_realidade_local.get("1.0", tk.END).strip(),
            "res_lousa_caderno": self.var_res_lousa_caderno.get(),
            "res_sem_impressao": self.var_res_sem_impressao.get(),
            "res_datashow": self.var_res_datashow.get(),
            "res_patio": self.var_res_patio.get(),
            "res_laboratorio": self.var_res_laboratorio.get()
        }
        try:
            with open(STATE_FILE, "w", encoding="utf-8") as f:
                json.dump(data, f, ensure_ascii=False, indent=4)
            messagebox.showinfo("Sucesso", "Rascunho salvo com sucesso! Tudo o que você preencheu foi guardado.")
        except Exception as e:
            messagebox.showerror("Erro", f"Não foi possível salvar o rascunho: {e}")

    def load_state(self):
        if os.path.exists(STATE_FILE):
            try:
                with open(STATE_FILE, "r", encoding="utf-8") as f:
                    data = json.load(f)

                if "file_esqueleto" in data: self.file_esqueleto_path.set(data["file_esqueleto"])
                if "esqueleto" in data:
                    self.txt_esqueleto.delete("1.0", tk.END)
                    self.txt_esqueleto.insert("1.0", data["esqueleto"])

                if "file_historico" in data: self.file_historico_path.set(data["file_historico"])
                if "historico" in data:
                    self.txt_historico.delete("1.0", tk.END)
                    self.txt_historico.insert("1.0", data["historico"])

                if "file_livro" in data: self.file_livro_path.set(data["file_livro"])
                if "livro" in data:
                    self.txt_livro.delete("1.0", tk.END)
                    self.txt_livro.insert("1.0", data["livro"])

                if "file_plano_base" in data: self.file_plano_base_path.set(data["file_plano_base"])

                if hasattr(self, 'txt_instrucoes_atividade') and "instrucoes_atividade" in data:
                    self.txt_instrucoes_atividade.delete("1.0", tk.END)
                    self.txt_instrucoes_atividade.insert("1.0", data["instrucoes_atividade"])

                if "genero" in data:
                    self.ent_genero.delete(0, tk.END)
                    self.ent_genero.insert(0, data["genero"])

                if "extra" in data:
                    self.ent_extra.delete(0, tk.END)
                    self.ent_extra.insert(0, data["extra"])

                if "frase" in data and data["frase"] in self.frases_list:
                    self.combo_frases.set(data["frase"])

                if "observacoes" in data:
                    self.txt_observacoes.delete("1.0", tk.END)
                    self.txt_observacoes.insert("1.0", data["observacoes"])

                if "ano" in data and data["ano"] in self.combo_ano['values']: self.combo_ano.set(data["ano"])
                if "duracao" in data: self.combo_duracao.set(data["duracao"])
                if "multi" in data: self.var_multi.set(data["multi"])
                if "auto_bncc" in data: self.var_auto_bncc.set(data["auto_bncc"])

                if "bncc_manual" in data:
                    self.ent_bncc_manual.delete(0, tk.END)
                    self.ent_bncc_manual.insert(0, data["bncc_manual"])

                if "escola_nome" in data:
                    self.ent_escola_nome.delete(0, tk.END)
                    self.ent_escola_nome.insert(0, data["escola_nome"])

                if "perfil_turma" in data:
                    self.txt_perfil_turma.delete("1.0", tk.END)
                    self.txt_perfil_turma.insert("1.0", data["perfil_turma"])

                if "infraestrutura" in data:
                    self.txt_infraestrutura.delete("1.0", tk.END)
                    self.txt_infraestrutura.insert("1.0", data["infraestrutura"])

                if "realidade_local" in data:
                    self.txt_realidade_local.delete("1.0", tk.END)
                    self.txt_realidade_local.insert("1.0", data["realidade_local"])

                if "res_lousa_caderno" in data: self.var_res_lousa_caderno.set(data["res_lousa_caderno"])
                if "res_sem_impressao" in data: self.var_res_sem_impressao.set(data["res_sem_impressao"])
                if "res_datashow" in data: self.var_res_datashow.set(data["res_datashow"])
                if "res_patio" in data: self.var_res_patio.set(data["res_patio"])
                if "res_laboratorio" in data: self.var_res_laboratorio.set(data["res_laboratorio"])

            except Exception as e:
                print(f"Aviso ao carregar rascunho anterior: {e}")

    def create_tabs(self):
        self.notebook = ttk.Notebook(self.root)
        self.notebook.pack(fill="both", expand=True, padx=10, pady=10)

        self.tab1 = ttk.Frame(self.notebook)
        self.tab2 = ttk.Frame(self.notebook)
        self.tab3 = ttk.Frame(self.notebook)
        self.tab4 = ttk.Frame(self.notebook)
        self.tab5 = ttk.Frame(self.notebook)
        self.tab6 = ttk.Frame(self.notebook)
        self.tab7 = ttk.Frame(self.notebook)

        self.notebook.add(self.tab1, text=" 📝 1. Esqueleto Oficial ")
        self.notebook.add(self.tab2, text=" 📚 2. Matriz & Histórico ")
        self.notebook.add(self.tab3, text=" 📖 3. Livro Didático/PDF ")
        self.notebook.add(self.tab4, text=" 💡 4. Conteúdo Extra & Frases ")
        self.notebook.add(self.tab5, text=" 🎯 5. BNCC Inteligente ")
        self.notebook.add(self.tab6, text=" 🏫 6. Contexto Escolar ")
        self.notebook.add(self.tab7, text=" 🖨️ 7. Gerar Atividades ")

        self.setup_tab1()
        self.setup_tab2()
        self.setup_tab3()
        self.setup_tab4()
        self.setup_tab5()
        self.setup_tab6()
        self.setup_tab7()

    def create_field_toolbar(self, parent, target_widget):
        bar = ttk.Frame(parent)
        bar.pack(anchor="e", pady=(0, 2))
        ttk.Button(bar, text="📋 Colar", width=8, command=lambda: paste_clipboard_to_widget(target_widget)).pack(side="left", padx=2)
        ttk.Button(bar, text="❌ Limpar", width=8, command=lambda: clear_widget(target_widget)).pack(side="left", padx=2)

    def setup_tab1(self):
        frame = ttk.LabelFrame(self.tab1, text=" Modelo Estrutural da Escola (.docx / .txt) ", padding=15)
        frame.pack(fill="both", expand=True, padx=15, pady=15)

        ttk.Label(frame, text="Selecione o arquivo do Word ou TXT que contém a estrutura oficial exigida pela sua escola:").pack(anchor="w", pady=(0, 5))

        file_box = ttk.Frame(frame)
        file_box.pack(fill="x", pady=5)
        ent = ttk.Entry(file_box, textvariable=self.file_esqueleto_path, width=60)
        ent.pack(side="left", fill="x", expand=True, padx=(0, 5))
        add_context_menu(ent)
        ttk.Button(file_box, text="Procurar...", command=lambda: self.pick_file(self.file_esqueleto_path, [("Word / Text", "*.docx *.txt")])).pack(side="right")

        ttk.Label(frame, text="Ou digite/cole a estrutura do plano diretamente no campo abaixo:").pack(anchor="w", pady=(15, 5))
        self.txt_esqueleto = scrolledtext.ScrolledText(frame, height=12, font=("Consolas", 10))
        self.txt_esqueleto.pack(fill="both", expand=True)
        add_context_menu(self.txt_esqueleto)
        self.create_field_toolbar(frame, self.txt_esqueleto)

    def setup_tab2(self):
        frame = ttk.LabelFrame(self.tab2, text=" Histórico Pedagógico / Planos Anteriores ", padding=15)
        frame.pack(fill="both", expand=True, padx=15, pady=15)

        ttk.Label(frame, text="Carregue planos de aula das semanas passadas para manter o estilo didático e a sequência cronológica:").pack(anchor="w", pady=(0, 5))

        file_box = ttk.Frame(frame)
        file_box.pack(fill="x", pady=5)
        ent = ttk.Entry(file_box, textvariable=self.file_historico_path, width=60)
        ent.pack(side="left", fill="x", expand=True, padx=(0, 5))
        add_context_menu(ent)
        ttk.Button(file_box, text="Procurar...", command=lambda: self.pick_file(self.file_historico_path, [("Word / Text", "*.docx *.txt")])).pack(side="right")

        ttk.Label(frame, text="Resumo de conteúdos ministrados recentemente:").pack(anchor="w", pady=(15, 5))
        self.txt_historico = scrolledtext.ScrolledText(frame, height=12, font=("Consolas", 10))
        self.txt_historico.pack(fill="both", expand=True)
        add_context_menu(self.txt_historico)
        self.create_field_toolbar(frame, self.txt_historico)

    def setup_tab3(self):
        frame = ttk.LabelFrame(self.tab3, text=" Ingestão do Material da Semana (Livro / Apostila) ", padding=15)
        frame.pack(fill="both", expand=True, padx=15, pady=15)

        ttk.Label(frame, text="Selecione o arquivo em PDF ou DOCX contendo as páginas do livro didático da semana:").pack(anchor="w", pady=(0, 5))

        file_box = ttk.Frame(frame)
        file_box.pack(fill="x", pady=5)
        ent = ttk.Entry(file_box, textvariable=self.file_livro_path, width=60)
        ent.pack(side="left", fill="x", expand=True, padx=(0, 5))
        add_context_menu(ent)
        ttk.Button(file_box, text="Procurar PDF/DOCX...", command=lambda: self.pick_file(self.file_livro_path, [("Documentos", "*.pdf *.docx *.txt")])).pack(side="right")

        ttk.Label(frame, text="Anotações / Trechos copiados do livro (Opcional):").pack(anchor="w", pady=(15, 5))
        self.txt_livro = scrolledtext.ScrolledText(frame, height=12, font=("Consolas", 10))
        self.txt_livro.pack(fill="both", expand=True)
        add_context_menu(self.txt_livro)
        self.create_field_toolbar(frame, self.txt_livro)

    def setup_tab4(self):
        frame = ttk.LabelFrame(self.tab4, text=" Conteúdos Complementares & Diretrizes Didáticas ", padding=15)
        frame.pack(fill="both", expand=True, padx=15, pady=15)

        ttk.Label(frame, text="Gênero Textual da Semana (Ex: Tirinha, Fábula, Poema, Conto, Artigo de Opinião):", font=("Segoe UI", 9, "bold")).pack(anchor="w", pady=(0, 2))
        self.ent_genero = ttk.Entry(frame, width=50)
        self.ent_genero.pack(anchor="w", pady=(0, 10))
        add_context_menu(self.ent_genero)

        ttk.Label(frame, text="Conteúdos Extra-Livro / Temas Complementares da Semana:", font=("Segoe UI", 9, "bold")).pack(anchor="w", pady=(0, 2))
        self.ent_extra = ttk.Entry(frame, width=70)
        self.ent_extra.pack(anchor="w", pady=(0, 10))
        add_context_menu(self.ent_extra)

        ttk.Label(frame, text="Frase Pedagógica Pronta (Ênfase Metodológica Selecionável):", font=("Segoe UI", 9, "bold")).pack(anchor="w", pady=(0, 2))
        self.frases_list = [
            "Focar em Atividades Práticas, Projetos e Aprendizagem Mão na Massa (Maker).",
            "Priorizar Exposição Teórica Interativa + Exercícios de Fixação e Sistematização.",
            "Enfatizar Leitura Crítica, Interpretação Textual e Rodas de Conversa/Debates.",
            "Aplicações Gamificadas, Jogos Pedagógicos em Duplas e Dinâmicas de Grupo.",
            "Foco em Consolidação de Aprendizagem, Diagnóstico e Acompanhamento Individualizado."
        ]
        self.combo_frases = ttk.Combobox(frame, values=self.frases_list, width=75, state="readonly")
        self.combo_frases.current(0)
        self.combo_frases.pack(anchor="w", pady=(0, 15))

        ttk.Label(frame, text="Observações e Orientações Específicas para esta Semana:").pack(anchor="w", pady=(0, 2))
        self.txt_observacoes = scrolledtext.ScrolledText(frame, height=6, font=("Segoe UI", 10))
        self.txt_observacoes.pack(fill="both", expand=True)
        add_context_menu(self.txt_observacoes)
        self.create_field_toolbar(frame, self.txt_observacoes)

    def setup_tab5(self):
        frame = ttk.LabelFrame(self.tab5, text=" Filtros da BNCC e Adequação Cognitiva (EF1, EF2 e EM) ", padding=15)
        frame.pack(fill="both", expand=True, padx=15, pady=15)

        box1 = ttk.Frame(frame)
        box1.pack(fill="x", pady=5)

        ttk.Label(box1, text="Série / Ano Escolar:", font=("Segoe UI", 10, "bold")).pack(side="left", padx=(0, 10))
        self.combo_ano = ttk.Combobox(box1, values=[
            "1º Ano", "2º Ano", "3º Ano", "4º Ano", "5º Ano",
            "6º Ano", "7º Ano", "8º Ano", "9º Ano",
            "1º Ano EM", "2º Ano EM", "3º Ano EM"
        ], state="readonly", width=14)
        self.combo_ano.current(4)
        self.combo_ano.pack(side="left", padx=(0, 20))

        ttk.Label(box1, text="Tempo / Duração:", font=("Segoe UI", 10, "bold")).pack(side="left", padx=(0, 10))
        self.combo_duracao = ttk.Combobox(box1, values=[
            "1 Aula (50 minutos)",
            "Bloco Duplo (100 minutos)",
            "Semanal Completo (5 Dias de Aula)",
            "Período Integral"
        ], state="readonly", width=25)
        self.combo_duracao.current(2)
        self.combo_duracao.pack(side="left")

        box2 = ttk.Frame(frame)
        box2.pack(fill="x", pady=10)
        self.var_multi = tk.BooleanVar(value=True)
        ttk.Checkbutton(box2, text="Plano Semanal Multidisciplinar / Integrado", variable=self.var_multi).pack(anchor="w")

        self.var_auto_bncc = tk.BooleanVar(value=True)
        ttk.Checkbutton(box2, text="Filtro Inteligente de Habilidades (O programa lerá o bncc_data.json e extrairá códigos válidos)", variable=self.var_auto_bncc).pack(anchor="w", pady=5)

        ttk.Label(frame, text="Inserir Códigos Específicos da BNCC Manuais (Opcional - Ex: EF05LP01, EM13LGG101):").pack(anchor="w", pady=(10, 2))
        self.ent_bncc_manual = ttk.Entry(frame, width=60)
        self.ent_bncc_manual.pack(anchor="w")
        add_context_menu(self.ent_bncc_manual)

        info_box = tk.Text(frame, height=6, bg="#F1F5F9", relief="flat", font=("Segoe UI", 9))
        info_box.insert("1.0", "🔒 Suporte Expandido BNCC (Educação Infantil até Ensino Médio):\nO EduPlan AI buscará automaticamente o arquivo 'bncc_data.json' embutido no sistema. Ele enviará à IA apenas as habilidades compatíveis com a série selecionada (incluindo o prefixo EM13 para o Ensino Médio).")
        info_box.configure(state="disabled")
        info_box.pack(fill="x", pady=(15, 0))

    def setup_tab6(self):
        frame = ttk.LabelFrame(self.tab6, text=" Realidade Local, Perfil da Turma, Recursos & Acessibilidade ", padding=15)
        frame.pack(fill="both", expand=True, padx=15, pady=15)

        ttk.Label(frame, text="Nome da Escola / Identificação da Turma (Opcional):", font=("Segoe UI", 9, "bold")).pack(anchor="w", pady=(0, 2))
        self.ent_escola_nome = ttk.Entry(frame, width=60)
        self.ent_escola_nome.pack(anchor="w", pady=(0, 10))
        add_context_menu(self.ent_escola_nome)

        ttk.Label(frame, text="Perfil da Turma e Inclusão (Ex: Alunos PCD, TDAH, Autismo, Nível de Alfabetização, Ritmo):", font=("Segoe UI", 9, "bold")).pack(anchor="w", pady=(0, 2))
        self.txt_perfil_turma = scrolledtext.ScrolledText(frame, height=3, font=("Segoe UI", 10))
        self.txt_perfil_turma.pack(fill="x", pady=(0, 10))
        add_context_menu(self.txt_perfil_turma)

        ttk.Label(frame, text="Atalhos Rápidos de Infraestrutura & Restrições de Recursos:", font=("Segoe UI", 9, "bold")).pack(anchor="w", pady=(5, 2))

        preset_box = ttk.Frame(frame)
        preset_box.pack(fill="x", pady=(0, 10))

        self.var_res_lousa_caderno = tk.BooleanVar(value=True)
        self.var_res_sem_impressao = tk.BooleanVar(value=True)
        self.var_res_datashow = tk.BooleanVar(value=False)
        self.var_res_patio = tk.BooleanVar(value=False)
        self.var_res_laboratorio = tk.BooleanVar(value=False)

        col1 = ttk.Frame(preset_box)
        col1.pack(side="left", fill="y", padx=(0, 20))
        ttk.Checkbutton(col1, text="📌 Apenas Quadro / Giz e Caderno (Escola com Baixo Recurso)", variable=self.var_res_lousa_caderno).pack(anchor="w")
        ttk.Checkbutton(col1, text="🚫 Sem impressões para alunos (Uso do Livro ou Lousa)", variable=self.var_res_sem_impressao).pack(anchor="w")

        col2 = ttk.Frame(preset_box)
        col2.pack(side="left", fill="y")
        ttk.Checkbutton(col2, text="🖥️ Projetor / Datashow / Lousa Digital Disponível", variable=self.var_res_datashow).pack(anchor="w")
        ttk.Checkbutton(col2, text="🏀 Pátio / Quadra de Esportes Liberada", variable=self.var_res_patio).pack(anchor="w")
        ttk.Checkbutton(col2, text="💻 Laboratório de Informática / Tablets Disponíveis", variable=self.var_res_laboratorio).pack(anchor="w")

        ttk.Label(frame, text="Outros Recursos Específicos ou Materiais Disponíveis (Caixa Adaptável):", font=("Segoe UI", 9, "bold")).pack(anchor="w", pady=(5, 2))
        self.txt_infraestrutura = scrolledtext.ScrolledText(frame, height=2, font=("Segoe UI", 10))
        self.txt_infraestrutura.pack(fill="x", pady=(0, 10))
        add_context_menu(self.txt_infraestrutura)
        self.create_field_toolbar(frame, self.txt_infraestrutura)

        ttk.Label(frame, text="Realidade Socioespacial e Projetos Locais da Escola:", font=("Segoe UI", 9, "bold")).pack(anchor="w", pady=(0, 2))
        self.txt_realidade_local = scrolledtext.ScrolledText(frame, height=3, font=("Segoe UI", 10))
        self.txt_realidade_local.pack(fill="both", expand=True)

    def setup_tab7(self):
        frame = ttk.LabelFrame(self.tab7, text=" Gerador de Atividades Impressas (Folha do Aluno) ", padding=15)
        frame.pack(fill="both", expand=True, padx=15, pady=15)

        ttk.Label(frame, text="1. Selecione o Plano de Aula (.docx) base (Opcional se preencher as Instruções abaixo):").pack(anchor="w", pady=(0, 5))

        file_box = ttk.Frame(frame)
        file_box.pack(fill="x", pady=5)
        ent = ttk.Entry(file_box, textvariable=self.file_plano_base_path, width=60)
        ent.pack(side="left", fill="x", expand=True, padx=(0, 5))
        add_context_menu(ent)
        ttk.Button(file_box, text="Procurar...", command=lambda: self.pick_file(self.file_plano_base_path, [("Word Document", "*.docx")])).pack(side="right")

        ttk.Label(frame, text="2. Instruções Específicas para a Folha de Atividades (Opcional se tiver Plano Base):", font=("Segoe UI", 9, "bold")).pack(anchor="w", pady=(15, 5))
        ttk.Label(frame, text="Ex: 'Crie 5 questões de múltipla escolha sobre tabela de preços'").pack(anchor="w", pady=(0, 2))

        self.txt_instrucoes_atividade = scrolledtext.ScrolledText(frame, height=8, font=("Consolas", 10))
        self.txt_instrucoes_atividade.pack(fill="both", expand=True)
        add_context_menu(self.txt_instrucoes_atividade)
        self.create_field_toolbar(frame, self.txt_instrucoes_atividade)

        self.btn_gerar_atividade = tk.Button(
            frame,
            text="🖨️ GERAR CADERNO DE ATIVIDADES NO WORD",
            command=self.start_activity_generation_thread,
            bg="#7C3AED",
            fg="white",
            font=("Segoe UI", 10, "bold"),
            padx=12,
            pady=8,
            relief="raised"
        )
        self.btn_gerar_atividade.pack(pady=10)

    def create_footer(self):
        footer_frame = tk.Frame(self.root)
        footer_frame.pack(fill="x", padx=10, pady=10)

        self.btn_gerar = tk.Button(
            footer_frame,
            text="🚀 GERAR PLANO DE AULA NO WORD (.DOCX)",
            command=self.start_generation_thread,
            bg="#2563EB",
            fg="white",
            font=("Segoe UI", 10, "bold"),
            padx=12,
            pady=8,
            relief="raised"
        )
        self.btn_gerar.pack(side="right", padx=4)

        self.btn_revisar = tk.Button(
            footer_frame,
            text="✏️ Revisão Cirúrgica",
            command=self.open_surgical_feedback_dialog,
            bg="#475569",
            fg="white",
            font=("Segoe UI", 10, "bold"),
            padx=10,
            pady=8,
            relief="raised"
        )
        self.btn_revisar.pack(side="right", padx=4)

        self.btn_salvar_rascunho = tk.Button(
            footer_frame,
            text="💾 Salvar Rascunho",
            command=self.save_state,
            bg="#059669",
            fg="white",
            font=("Segoe UI", 10, "bold"),
            padx=10,
            pady=8,
            relief="raised"
        )
        self.btn_salvar_rascunho.pack(side="right", padx=4)

        self.lbl_status = ttk.Label(footer_frame, text="Pronto para trabalhar. Motor Multi-IA Ativo.", font=("Segoe UI", 10))
        self.lbl_status.pack(side="left", padx=10)

    def pick_file(self, target_var, file_types):
        path = filedialog.askopenfilename(filetypes=file_types)
        if path:
            target_var.set(path)

    def extract_text_from_file(self, filepath):
        if not filepath or not os.path.exists(filepath):
            return ""
        ext = os.path.splitext(filepath)[1].lower()
        text = ""
        try:
            if ext == ".txt":
                with open(filepath, "r", encoding="utf-8", errors="ignore") as f:
                    text = f.read()
            elif ext == ".docx" and docx:
                doc = docx.Document(filepath)
                text = "\n".join([p.text for p in doc.paragraphs if p.text])
            elif ext == ".pdf" and pypdf:
                reader = pypdf.PdfReader(filepath)
                for page in reader.pages:
                    extracted = page.extract_text()
                    if extracted:
                        text += extracted + "\n"
        except Exception as e:
            print(f"Erro ao ler arquivo {filepath}: {e}")
        return text.strip()

    # ================= LOGICA DO PLANO DE AULA (MULTI-IA) =================
    def start_generation_thread(self):
        g_key = self.api_key_var.get().strip()
        gr_key = self.groq_key_var.get().strip()
        op_key = self.openrouter_key_var.get().strip()

        if not g_key and not gr_key and not op_key:
            messagebox.showerror("Erro", "Configure ao menos uma chave de API (Gemini, Groq ou OpenRouter) em Configurações.")
            self.open_settings_dialog()
            return

        self.btn_gerar.config(state="disabled", bg="#94A3B8")
        self.lbl_status.config(text="⏳ Processando BNCC e consultando motor de IA...")

        thread = threading.Thread(target=self.run_generation_process)
        thread.start()

    def run_generation_process(self):
        try:
            txt_esqueleto = self.extract_text_from_file(self.file_esqueleto_path.get()) + "\n" + self.txt_esqueleto.get("1.0", tk.END).strip()
            txt_historico = self.extract_text_from_file(self.file_historico_path.get()) + "\n" + self.txt_historico.get("1.0", tk.END).strip()
            txt_livro = self.extract_text_from_file(self.file_livro_path.get()) + "\n" + self.txt_livro.get("1.0", tk.END).strip()

            genero = self.ent_genero.get().strip()
            extra = self.ent_extra.get().strip()
            frase_diretriz = self.combo_frases.get()
            obs = self.txt_observacoes.get("1.0", tk.END).strip()

            ano_selecionado = self.combo_ano.get()
            duracao_selecionada = self.combo_duracao.get()
            is_multi = self.var_multi.get()
            bncc_manual = self.ent_bncc_manual.get().strip()

            escola_nome = self.ent_escola_nome.get().strip()
            perfil_turma = self.txt_perfil_turma.get("1.0", tk.END).strip()
            infraestrutura_texto = self.txt_infraestrutura.get("1.0", tk.END).strip()
            realidade_local = self.txt_realidade_local.get("1.0", tk.END).strip()

            infra_presets = []
            if self.var_res_lousa_caderno.get():
                infra_presets.append("APENAS LOUSA/GIZ E CADERNO (Escola de baixo recurso).")
            if self.var_res_sem_impressao.get():
                infra_presets.append("SEM IMPRESSÕES INDIVIDUAIS PARA ALUNOS.")
            if self.var_res_datashow.get():
                infra_presets.append("Datashow disponível.")
            if self.var_res_patio.get():
                infra_presets.append("Pátio liberado.")
            if self.var_res_laboratorio.get():
                infra_presets.append("Laboratório disponível.")

            infra_final = "; ".join(infra_presets)
            if infraestrutura_texto:
                infra_final += f" | Outros: {infraestrutura_texto}"

            habilidades_filtradas = []
            bncc_formatted_text = ""

            if os.path.exists(bncc_json_path) and self.var_auto_bncc.get():
                try:
                    with open(bncc_json_path, 'r', encoding='utf-8') as f:
                        bncc_data = json.load(f)

                    prefixos_bncc = {
                        "1º Ano": ["EF01", "EF12", "EF15"],
                        "2º Ano": ["EF02", "EF12", "EF15"],
                        "3º Ano": ["EF03", "EF35", "EF15"],
                        "4º Ano": ["EF04", "EF35", "EF15"],
                        "5º Ano": ["EF05", "EF35", "EF15"],
                        "6º Ano": ["EF06", "EF67", "EF69"],
                        "7º Ano": ["EF07", "EF67", "EF69"],
                        "8º Ano": ["EF08", "EF89", "EF69"],
                        "9º Ano": ["EF09", "EF89", "EF69"],
                        "1º Ano EM": ["EM13", "EM"],
                        "2º Ano EM": ["EM13", "EM"],
                        "3º Ano EM": ["EM13", "EM"]
                    }

                    prefixos_validos = prefixos_bncc.get(ano_selecionado, [])
                    for item in bncc_data:
                        codigo = item.get("codigo", "")
                        if any(codigo.startswith(prefixo) for prefixo in prefixos_validos):
                            habilidades_filtradas.append(item)

                    if habilidades_filtradas:
                        linhas_bncc = [f"[{h.get('codigo', '')}] {h.get('descricao', '').replace(chr(10), ' ')}" for h in habilidades_filtradas]
                        bncc_formatted_text = "\n".join(linhas_bncc)
                except Exception as e:
                    print(f"Erro ao processar JSON: {e}")

            system_instruction = (
                "Você é o motor pedagógico especialista do EduPlan AI.\n"
                "Responda DIRETAMENTE com o plano de aula em Markdown limpo. Sem saudações.\n"
                "REGRAS DE FORMATAÇÃO: NUNCA use formatação em LaTeX (como $\\frac{...}{...}$ ou $...$). Escreva frações como 1/2 e porcentagens como 25%. "
                "Não use o sinal > para blocos de citação ou texto de apoio. "
                "Respeite estritamente os códigos da BNCC fornecidos e as restrições de recursos."
            )

            prompt_user = f"""
- Série: {ano_selecionado} | Duração: {duracao_selecionada} | Multidisciplinar: {is_multi}
- Escola: {escola_nome} | Perfil da Turma: {perfil_turma} | Infra: {infra_final}
- Esqueleto: {txt_esqueleto}
- Histórico: {txt_historico}
- Livro: {txt_livro}
- Gênero: {genero} | Extra: {extra} | Ênfase: {frase_diretriz} | Obs: {obs}
- Habilidades BNCC Disponíveis:\n{bncc_formatted_text}
- Manuais: {bncc_manual}
Gere o plano completo.
"""
            resposta_markdown, telemetria = call_ai_multi_provider(
                gemini_key=self.api_key_var.get().strip(),
                groq_key=self.groq_key_var.get().strip(),
                openrouter_key=self.openrouter_key_var.get().strip(),
                provider_priority=self.provider_priority_var.get().strip(),
                prompt=prompt_user,
                system_instruction=system_instruction,
                temperature=0.6
            )

            target_dir = self.save_dir_var.get().strip()
            if not os.path.isdir(target_dir):
                target_dir = BASE_DIR

            timestamp = datetime.datetime.now().strftime("%Y-%m-%d_%H-%M-%S")
            output_filename = os.path.join(target_dir, f"Plano_de_Aula_{ano_selecionado.replace(' ', '_')}_{timestamp}.docx")

            self.export_markdown_to_docx(resposta_markdown, output_filename)
            self.file_plano_base_path.set(output_filename)

            msg_status = f"✅ Sucesso! [Provedor: {telemetria['model']} | Tokens: {telemetria['total_tokens']}]"
            self.lbl_status.config(text=msg_status)

            messagebox.showinfo("Sucesso", f"Plano de aula gerado com sucesso!\n\n📊 Telemetria da IA:\n• Modelo Utilizado: {telemetria['model']}\n• Tokens de Entrada: {telemetria['prompt_tokens']}\n• Tokens de Saída: {telemetria['response_tokens']}\n• Total de Tokens: {telemetria['total_tokens']}\n\nSalvo em:\n{output_filename}")

            if os.name == 'posix':
                os.system(f"xdg-open '{output_filename}'")
            else:
                os.startfile(output_filename)

        except Exception as e:
            self.lbl_status.config(text="❌ Erro na geração.")
            messagebox.showerror("Erro na Execução", f"Ocorreu um erro: {str(e)}")
        finally:
            self.btn_gerar.config(state="normal", bg="#2563EB")


    # ================= LOGICA DE ATIVIDADES (MULTI-IA) =================
    def start_activity_generation_thread(self):
        g_key = self.api_key_var.get().strip()
        gr_key = self.groq_key_var.get().strip()
        op_key = self.openrouter_key_var.get().strip()

        if not g_key and not gr_key and not op_key:
            messagebox.showerror("Erro", "Configure ao menos uma chave de API em Configurações.")
            return

        plano_path = self.file_plano_base_path.get().strip()
        instrucoes = self.txt_instrucoes_atividade.get("1.0", tk.END).strip()

        if (not plano_path or not os.path.exists(plano_path)) and not instrucoes:
            messagebox.showwarning("Aviso", "Forneça um Plano de Aula (.docx) OU digite Instruções Específicas na Aba 7 para gerar as atividades.")
            return

        self.btn_gerar_atividade.config(state="disabled", bg="#A78BFA")
        self.lbl_status.config(text="⏳ Gerando Atividades via Motor Multi-IA...")

        thread = threading.Thread(target=self.run_activity_generation_process, args=(plano_path,))
        thread.start()

    def run_activity_generation_process(self, plano_path):
        try:
            texto_plano = ""
            if plano_path and os.path.exists(plano_path):
                texto_plano = self.extract_text_from_file(plano_path)

            txt_livro = self.extract_text_from_file(self.file_livro_path.get()) + "\n" + self.txt_livro.get("1.0", tk.END).strip()
            perfil_turma = self.txt_perfil_turma.get("1.0", tk.END).strip()
            ano_selecionado = self.combo_ano.get()
            escola_nome = self.ent_escola_nome.get().strip()
            instrucoes_customizadas = self.txt_instrucoes_atividade.get("1.0", tk.END).strip()

            system_instruction = (
                "Você é um Designer Instrucional Pedagógico.\n"
                "Crie uma FOLHA DE ATIVIDADES PARA O ALUNO responder em Markdown limpo.\n"
                "REGRAS DE FORMATAÇÃO: NUNCA use formatação em LaTeX (como $\\frac{...}{...}$ ou $...$). Escreva frações como 1/2 e porcentagens como 25%. "
                "Não use o sinal > para blocos de citação ou texto de apoio. "
                "Para linhas de resposta, use pontilhados padronizados (..........) sem misturar com traços. "
                "Inclua cabeçalho padrão e questões claras. "
                "Você pode e deve usar tabelas formatadas em markdown (ex: | Col 1 | Col 2 |) se precisar organizar dados em linhas e colunas."
            )

            prompt_plano = f"- Plano Base: {texto_plano}\n" if texto_plano else "- Plano Base: [Nenhum documento anexado. Basear-se EXCLUSIVAMENTE nas Instruções Específicas e contexto fornecidos abaixo.]\n"

            prompt_user = f"""
- Série: {ano_selecionado} | Escola: {escola_nome} | Perfil: {perfil_turma}
{prompt_plano}- Material Didático: {txt_livro}
- Instruções Específicas: {instrucoes_customizadas}
Gere a folha de atividades formatada.
"""
            resposta_markdown, telemetria = call_ai_multi_provider(
                gemini_key=self.api_key_var.get().strip(),
                groq_key=self.groq_key_var.get().strip(),
                openrouter_key=self.openrouter_key_var.get().strip(),
                provider_priority=self.provider_priority_var.get().strip(),
                prompt=prompt_user,
                system_instruction=system_instruction,
                temperature=0.7
            )

            target_dir = self.save_dir_var.get().strip()
            if not os.path.isdir(target_dir):
                target_dir = BASE_DIR

            timestamp = datetime.datetime.now().strftime("%Y-%m-%d_%H-%M-%S")
            output_filename = os.path.join(target_dir, f"Atividades_Impressas_{ano_selecionado.replace(' ', '_')}_{timestamp}.docx")

            self.export_markdown_to_docx(resposta_markdown, output_filename)

            msg_status = f"✅ Atividades Geradas! [Provedor: {telemetria['model']} | Tokens: {telemetria['total_tokens']}]"
            self.lbl_status.config(text=msg_status)

            messagebox.showinfo("Sucesso", f"Folha de atividades gerada com sucesso!\n\n📊 Telemetria da IA:\n• Modelo Utilizado: {telemetria['model']}\n• Total de Tokens: {telemetria['total_tokens']}\n\nSalva em:\n{output_filename}")

            if os.name == 'posix':
                os.system(f"xdg-open '{output_filename}'")
            else:
                os.startfile(output_filename)

        except Exception as e:
            self.lbl_status.config(text="❌ Erro na geração de atividades.")
            messagebox.showerror("Erro", f"Erro: {str(e)}")
        finally:
            self.btn_gerar_atividade.config(state="normal", bg="#7C3AED")

    def open_surgical_feedback_dialog(self):
        target_dir = self.save_dir_var.get().strip()
        if not os.path.isdir(target_dir):
            target_dir = BASE_DIR

        docx_files = [f for f in os.listdir(target_dir) if f.endswith(".docx")]

        win = tk.Toplevel(self.root)
        win.title("EduPlan AI — Módulo de Feedback Cirúrgico")
        win.geometry("650x580")
        win.minsize(600, 500)

        tk.Label(win, text="Selecione o Documento gerado recentemente (.docx):", font=("Segoe UI", 10, "bold")).pack(anchor="w", padx=20, pady=(15, 5))

        file_combo = ttk.Combobox(win, values=docx_files, width=70, state="readonly")
        if docx_files:
            file_combo.current(0)
        file_combo.pack(anchor="w", padx=20, pady=(0, 10))

        tk.Label(win, text="Trecho ou Seção a ser Ajustada:", font=("Segoe UI", 10, "bold")).pack(anchor="w", padx=20, pady=(5, 5))

        txt_trecho = scrolledtext.ScrolledText(win, height=6, font=("Consolas", 9))
        txt_trecho.pack(fill="x", padx=20, pady=(0, 5))
        add_context_menu(txt_trecho)
        self.create_field_toolbar(win, txt_trecho)

        tk.Label(win, text="Instrução de Ajuste Cirúrgico:", font=("Segoe UI", 10, "bold")).pack(anchor="w", padx=20, pady=(5, 5))

        ent_ajuste = ttk.Entry(win, width=75)
        ent_ajuste.pack(anchor="w", padx=20, pady=(0, 15))
        add_context_menu(ent_ajuste)

        status_lbl = ttk.Label(win, text="Pronto para reescrever trecho.", font=("Segoe UI", 9))
        status_lbl.pack(anchor="w", padx=20, pady=(0, 10))

        def executar_revisao_pontual():
            selected_file = file_combo.get()
            trecho_alvo = txt_trecho.get("1.0", tk.END).strip()
            instrucao = ent_ajuste.get().strip()

            if not selected_file or not instrucao:
                messagebox.showerror("Erro", "Selecione um arquivo e informe a instrução.", parent=win)
                return

            filepath = os.path.join(target_dir, selected_file)
            texto_atual = self.extract_text_from_file(filepath)

            status_lbl.config(text="⏳ Aplicando revisão...")
            win.update()

            try:
                system_instruction = "Você é o editor pedagógico do EduPlan AI. Reescreva APENAS o trecho solicitado mantendo a coerência. REGRAS: Não use LaTeX ou sinal de maior (>) para citações."
                prompt_surgical = f"[DOCUMENTO]\n{texto_atual}\n\n[TRECHO ALVO]\n{trecho_alvo}\n\n[AJUSTE]\n{instrucao}"

                novo_texto, telemetria = call_ai_multi_provider(
                    gemini_key=self.api_key_var.get().strip(),
                    groq_key=self.groq_key_var.get().strip(),
                    openrouter_key=self.openrouter_key_var.get().strip(),
                    provider_priority=self.provider_priority_var.get().strip(),
                    prompt=prompt_surgical,
                    system_instruction=system_instruction,
                    temperature=0.5
                )

                base_name, ext = os.path.splitext(filepath)
                timestamp = datetime.datetime.now().strftime("%Y-%m-%d_%H-%M-%S")
                novo_filepath = f"{base_name}_revisado_{timestamp}{ext}"

                self.export_markdown_to_docx(novo_texto, novo_filepath)

                status_lbl.config(text=f"✅ OK! [Provedor: {telemetria['model']} | Tokens: {telemetria['total_tokens']}]")
                messagebox.showinfo("Sucesso", f"Documento revisado salvo em:\n{novo_filepath}", parent=win)

                if os.name == 'posix':
                    os.system(f"xdg-open '{novo_filepath}'")
                else:
                    os.startfile(novo_filepath)
                win.destroy()

            except Exception as e:
                status_lbl.config(text="❌ Erro na revisão.")
                messagebox.showerror("Erro", f"Falha: {str(e)}", parent=win)

        tk.Button(win, text="⚡ Executar Revisão Cirúrgica", command=executar_revisao_pontual, bg="#2563EB", fg="white", font=("Segoe UI", 10, "bold"), padx=15, pady=5).pack(pady=5)

    def render_table_to_docx(self, doc, table_lines):
        parsed_rows = []
        for t_line in table_lines:
            clean_line = t_line.replace('|', '').replace('-', '').replace(':', '').strip()
            if clean_line == '':
                continue

            cells = [c.strip() for c in t_line.split('|')]
            if t_line.startswith('|'): cells = cells[1:]
            if t_line.endswith('|') and len(cells) > 0: cells = cells[:-1]

            parsed_rows.append(cells)

        if not parsed_rows:
            return

        cols = max(len(r) for r in parsed_rows) if parsed_rows else 0
        if cols == 0: return

        table = doc.add_table(rows=len(parsed_rows), cols=cols)

        try:
            table.style = 'Table Grid'
        except:
            pass

        for row_idx, row_data in enumerate(parsed_rows):
            for col_idx, cell_text in enumerate(row_data):
                if col_idx < len(table.columns):
                    cell = table.cell(row_idx, col_idx)
                    p = cell.paragraphs[0]
                    p.paragraph_format.space_after = Pt(2)

                    if row_idx == 0:
                        run = p.add_run(cell_text.replace("**", "").replace("*", ""))
                        run.font.bold = True
                        run.font.name = "Arial"
                        run.font.size = Pt(10)
                    else:
                        self.add_formatted_text(p, cell_text)

        doc.add_paragraph()

    def export_markdown_to_docx(self, md_text, filename):
        if not docx:
            raise ImportError("A biblioteca python-docx não está instalada.")

        doc = docx.Document()

        for section in doc.sections:
            section.top_margin = Inches(0.8)
            section.bottom_margin = Inches(0.8)
            section.left_margin = Inches(0.8)
            section.right_margin = Inches(0.8)

        lines = md_text.split("\n")
        table_lines_buffer = []

        for line in lines:
            line_str = line.strip()

            if line_str.startswith(">"):
                line_str = line_str.lstrip("> ").strip()

            line_str = re.sub(r'\$\\frac\{(\d+)\}\{(\d+)\}\$', r'\1/\2', line_str)
            line_str = re.sub(r'\$([\d.,]+%)\$', r'\1', line_str)
            line_str = line_str.replace("\\%", "%")

            if not line_str:
                if table_lines_buffer:
                    self.render_table_to_docx(doc, table_lines_buffer)
                    table_lines_buffer = []
                continue

            is_table_line = line_str.startswith("|") and line_str.endswith("|")
            if is_table_line:
                table_lines_buffer.append(line_str)
                continue
            else:
                if table_lines_buffer:
                    self.render_table_to_docx(doc, table_lines_buffer)
                    table_lines_buffer = []

            if line_str.startswith("#### "):
                p = doc.add_paragraph()
                run = p.add_run(line_str[5:].replace("**", ""))
                run.font.name = "Arial"
                run.font.size = Pt(11)
                run.font.bold = True
                run.font.color.rgb = RGBColor(0x1E, 0x3A, 0x8A)
                p.paragraph_format.space_before = Pt(8)
                p.paragraph_format.space_after = Pt(2)

            elif line_str.startswith("# "):
                p = doc.add_paragraph()
                run = p.add_run(line_str[2:].replace("**", ""))
                run.font.name = "Arial"
                run.font.size = Pt(14)
                run.font.bold = True
                run.font.color.rgb = RGBColor(0x1E, 0x3A, 0x8A)
                p.paragraph_format.space_before = Pt(10)
                p.paragraph_format.space_after = Pt(4)

            elif line_str.startswith("## "):
                p = doc.add_paragraph()
                run = p.add_run(line_str[3:].replace("**", ""))
                run.font.name = "Arial"
                run.font.size = Pt(12)
                run.font.bold = True
                run.font.color.rgb = RGBColor(0x25, 0x63, 0xEB)
                p.paragraph_format.space_before = Pt(8)
                p.paragraph_format.space_after = Pt(3)

            elif line_str.startswith("### "):
                p = doc.add_paragraph()
                run = p.add_run(line_str[4:].replace("**", ""))
                run.font.name = "Arial"
                run.font.size = Pt(11)
                run.font.bold = True
                run.font.color.rgb = RGBColor(0x33, 0x41, 0x55)
                p.paragraph_format.space_before = Pt(8)
                p.paragraph_format.space_after = Pt(2)

            elif line_str.startswith("- ") or line_str.startswith("* "):
                p = doc.add_paragraph(style='List Bullet')
                self.add_formatted_text(p, line_str[2:])
                p.paragraph_format.space_after = Pt(2)

            else:
                p = doc.add_paragraph()
                self.add_formatted_text(p, line_str)
                p.paragraph_format.space_after = Pt(3)

        if table_lines_buffer:
            self.render_table_to_docx(doc, table_lines_buffer)

        doc.save(filename)

    def add_formatted_text(self, paragraph, text):
        text = re.sub(r'\.{2,}', '________________________________________', text)

        parts = re.split(r'(\*\*.*?\*\*)', text)
        for part in parts:
            if part.startswith("**") and part.endswith("**"):
                clean_part = part[2:-2].replace("*", "")
                run = paragraph.add_run(clean_part)
                run.font.name = "Arial"
                run.font.size = Pt(10)
                run.font.bold = True
            else:
                clean_part = part.replace("*", "")
                run = paragraph.add_run(clean_part)
                run.font.name = "Arial"
                run.font.size = Pt(10)

if __name__ == "__main__":
    root = tk.Tk()
    app = EduPlanAIApp(root)
    root.mainloop()
